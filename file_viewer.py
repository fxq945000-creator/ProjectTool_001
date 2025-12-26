import os
import tkinter as tk
import tkinter.font as tkfont
from tkinter import ttk, filedialog, messagebox, simpledialog

import json
import csv

import time
import datetime
import sys
import subprocess

import shutil
import threading
import winreg

class ProgressDialog:
    def __init__(self, parent, title="处理中"):
        self.cancelled = False
        self.top = tk.Toplevel(parent)
        self.top.title(title)
        self.top.geometry("450x150")
        self.top.resizable(False, False)
        
        # 居中计算
        try:
            parent_x = parent.winfo_rootx()
            parent_y = parent.winfo_rooty()
            parent_w = parent.winfo_width()
            parent_h = parent.winfo_height()
            x = parent_x + (parent_w - 450) // 2
            y = parent_y + (parent_h - 150) // 2
            self.top.geometry(f"+{x}+{y}")
        except:
            # Fallback center on screen
            screen_width = parent.winfo_screenwidth()
            screen_height = parent.winfo_screenheight()
            x = (screen_width - 450) // 2
            y = (screen_height - 150) // 2
            self.top.geometry(f"+{x}+{y}")
            
        self.message_var = tk.StringVar(value="准备中...")
        self.progress_var = tk.DoubleVar(value=0)
        
        frame = ttk.Frame(self.top, padding=20)
        frame.pack(fill="both", expand=True)
        
        tk.Label(frame, textvariable=self.message_var, wraplength=410, justify="left", anchor="w").pack(pady=(0, 10), fill="x")
        self.pb = ttk.Progressbar(frame, variable=self.progress_var, maximum=100)
        self.pb.pack(fill="x", pady=5)
        
        ttk.Button(frame, text="取消", command=self.cancel).pack(pady=5)
        
        self.top.protocol("WM_DELETE_WINDOW", self.cancel)
        self.top.attributes("-topmost", True) # 强制置顶
        
        # 强制聚焦和刷新
        self.top.deiconify()
        self.top.lift()
        self.top.focus_force()
        self.top.update()
        
    def cancel(self):
        self.cancelled = True
        self.message_var.set("正在取消...")

    def update(self, percent, message=None):
        try:
            if not self.top.winfo_exists(): return
            self.progress_var.set(percent)
            if message:
                self.message_var.set(message)
            self.top.update_idletasks()
        except:
            pass
        
    def close(self):
        try:
            self.top.destroy()
        except:
            pass

class CreateToolTip(object):
    """
    create a tooltip for a given widget
    """
    def __init__(self, widget, text='widget info'):
        self.waittime = 500     # miliseconds
        self.wraplength = 180   # pixels
        self.widget = widget
        self.text = text
        self.widget.bind("<Enter>", self.enter)
        self.widget.bind("<Leave>", self.leave)
        self.widget.bind("<ButtonPress>", self.leave)
        self.id = None
        self.tw = None

    def enter(self, event=None):
        self.schedule()

    def leave(self, event=None):
        self.unschedule()
        self.hidetip()

    def schedule(self):
        self.unschedule()
        self.id = self.widget.after(self.waittime, self.showtip)

    def unschedule(self):
        id = self.id
        self.id = None
        if id:
            self.widget.after_cancel(id)

    def showtip(self, event=None):
        x = y = 0
        x, y, cx, cy = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 20
        # creates a toplevel window
        self.tw = tk.Toplevel(self.widget)
        # Leaves only the label and removes the app window
        self.tw.wm_overrideredirect(True)
        self.tw.wm_geometry("+%d+%d" % (x, y))
        label = tk.Label(self.tw, text=self.text, justify='left',
                       background="#ffffe0", relief='solid', borderwidth=1,
                       wraplength = self.wraplength)
        label.pack(ipadx=1)

    def hidetip(self):
        tw = self.tw
        self.tw= None
        if tw:
            tw.destroy()

try:
    from PIL import Image, ImageTk
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False

# 延迟导入 cv2 以防止启动卡顿
HAS_CV2 = False
try:
    import cv2
    HAS_CV2 = True
except ImportError:
    pass

# 支持预览的文本文件后缀
TEXT_EXTS = {
    ".txt", ".md", ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg",
    ".xml", ".html", ".htm", ".css", ".scss", ".less", ".sh", ".bat", ".ps1", ".go", ".rs", ".java", ".kt",
    ".c", ".cc", ".cpp", ".h", ".hpp", ".cs", ".sql", ".log", ".vue", ".lua", ".rb", ".php"
}

# 支持预览的图片文件后缀
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".ico", ".webp", ".tif", ".tiff"}

# 支持预览的视频文件后缀
VIDEO_EXTS = {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv", ".webm", ".m4v", ".ts"}

MAX_PREVIEW_BYTES = 2 * 1024 * 1024  # 2MB 限制

# === 版本与更新配置 ===
APP_VERSION = "1.0.1"
# [配置说明] 请将以下 URL 替换为您实际存放 version.txt 和最新代码的地址
# version.txt 内容示例: 1.0.2
UPDATE_CHECK_URL = "" 
UPDATE_DOWNLOAD_URL = ""

def is_text_file(path: str) -> bool:
    """简单判断是否为文本文件"""
    if os.path.isdir(path):
        return False
    _, ext = os.path.splitext(path)
    return ext.lower() in TEXT_EXTS

def is_image_file(path: str) -> bool:
    """简单判断是否为图片文件"""
    if os.path.isdir(path):
        return False
    _, ext = os.path.splitext(path)
    return ext.lower() in IMAGE_EXTS

def is_video_file(path: str) -> bool:
    """简单判断是否为视频文件"""
    if os.path.isdir(path):
        return False
    _, ext = os.path.splitext(path)
    return ext.lower() in VIDEO_EXTS

def read_file_content(path: str, max_bytes: int = MAX_PREVIEW_BYTES) -> str:
    """安全读取文件内容"""
    try:
        size = os.path.getsize(path)
        if size > max_bytes:
            return f"[系统提示] 文件过大 ({size} bytes)，已跳过预览以防卡顿。\n文件路径: {path}"
        
        with open(path, "rb") as f:
            raw = f.read(max_bytes + 1)
        
        # 简单的二进制检测
        if b"\x00" in raw:
            return f"[系统提示] 检测到二进制内容，无法以文本形式预览。\n文件路径: {path}"
            
        # 尝试解码
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return raw.decode("gbk")
            except UnicodeDecodeError:
                return raw.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[读取错误] {e}\n文件路径: {path}"

# === Windows Clipboard Utils ===
try:
    import ctypes
    from ctypes import wintypes
    
    # 64-bit safe types
    if ctypes.sizeof(ctypes.c_void_p) == 8:
        ULONG_PTR = ctypes.c_uint64
        LRESULT = ctypes.c_int64
    else:
        ULONG_PTR = ctypes.c_ulong
        LRESULT = ctypes.c_long

    # API Definitions
    _GetClipboardData = ctypes.windll.user32.GetClipboardData
    _GetClipboardData.argtypes = [ctypes.c_uint]
    _GetClipboardData.restype = ctypes.c_void_p # Handle is void*

    _OpenClipboard = ctypes.windll.user32.OpenClipboard
    _OpenClipboard.argtypes = [ctypes.c_void_p]
    _OpenClipboard.restype = ctypes.c_bool

    _CloseClipboard = ctypes.windll.user32.CloseClipboard
    _CloseClipboard.argtypes = []
    _CloseClipboard.restype = ctypes.c_bool
    
    _GlobalAlloc = ctypes.windll.kernel32.GlobalAlloc
    _GlobalAlloc.argtypes = [ctypes.c_uint, ctypes.c_size_t]
    _GlobalAlloc.restype = ctypes.c_void_p

    _GlobalLock = ctypes.windll.kernel32.GlobalLock
    _GlobalLock.argtypes = [ctypes.c_void_p]
    _GlobalLock.restype = ctypes.c_void_p

    _GlobalUnlock = ctypes.windll.kernel32.GlobalUnlock
    _GlobalUnlock.argtypes = [ctypes.c_void_p]
    _GlobalUnlock.restype = ctypes.c_bool
    
    _DragQueryFileW = ctypes.windll.shell32.DragQueryFileW
    _DragQueryFileW.argtypes = [ctypes.c_void_p, ctypes.c_uint, ctypes.c_wchar_p, ctypes.c_uint]
    _DragQueryFileW.restype = ctypes.c_uint
    
    _DragQueryPoint = ctypes.windll.shell32.DragQueryPoint
    _DragQueryPoint.argtypes = [ctypes.c_void_p, ctypes.POINTER(wintypes.POINT)]
    _DragQueryPoint.restype = ctypes.c_bool

    class ClipboardUtils:
        @staticmethod
        def set_files(paths):
            try:
                # Structure for DROPFILES
                class DROPFILES(ctypes.Structure):
                    _fields_ = [("pFiles", wintypes.DWORD),
                                ("pt", wintypes.POINT),
                                ("fNC", wintypes.BOOL),
                                ("fWide", wintypes.BOOL)]
                
                # Calculate size
                files_text = "\0".join(paths) + "\0\0"
                files_data = files_text.encode("utf-16le")
                dropfiles_size = ctypes.sizeof(DROPFILES)
                total_size = dropfiles_size + len(files_data)
                
                # Allocate global memory
                hGlobal = _GlobalAlloc(0x0042, total_size) # GHND
                if not hGlobal: return
                
                # Lock and Write
                ptr = _GlobalLock(hGlobal)
                if not ptr: 
                    ctypes.windll.kernel32.GlobalFree(hGlobal)
                    return
                
                # Write DROPFILES struct
                # We need to write memory directly
                df = DROPFILES()
                df.pFiles = dropfiles_size
                df.fWide = True
                
                ctypes.memmove(ptr, ctypes.byref(df), dropfiles_size)
                ctypes.memmove(ptr + dropfiles_size, files_data, len(files_data))
                _GlobalUnlock(hGlobal)
                
                # Set Clipboard
                if _OpenClipboard(None):
                    ctypes.windll.user32.EmptyClipboard()
                    ctypes.windll.user32.SetClipboardData(15, hGlobal) # CF_HDROP
                    _CloseClipboard()
            except Exception as e:
                print(f"Clipboard set error: {e}")

        @staticmethod
        def get_files():
            files = []
            try:
                # 尝试打开剪贴板，如果失败重试一次
                success = False
                for _ in range(3):
                    if _OpenClipboard(None):
                        success = True
                        break
                    time.sleep(0.1)
                    
                if not success:
                    return []

                hDrop = _GetClipboardData(15) # CF_HDROP
                if hDrop:
                    count = _DragQueryFileW(hDrop, 0xFFFFFFFF, None, 0)
                    buf = ctypes.create_unicode_buffer(4096)
                    for i in range(count):
                        length = _DragQueryFileW(hDrop, i, None, 0)
                        if length > 4096:
                             # 极少数情况
                             pass
                        else:
                            if _DragQueryFileW(hDrop, i, buf, 4096):
                                files.append(buf.value)
                _CloseClipboard()
            except Exception as e:
                print(f"Clipboard get error: {e}")
            return files

    # === Windows DnD Hook ===
    # 使用全局队列解耦 C 回调和 Python 主线程，防止 GIL 冲突
    import queue
    _dnd_queue = queue.Queue()

    WNDPROC = ctypes.WINFUNCTYPE(LRESULT, ctypes.c_void_p, ctypes.c_uint, ctypes.c_void_p, ctypes.c_void_p)
    
    class WindowsDnD:
        def __init__(self, widget, on_drop_callback):
            self.widget = widget
            self.on_drop_callback = on_drop_callback
            self.old_wnd_proc = None
            self.new_wnd_proc = None
            self.hwnd = None

        def hook(self):
            try:
                # 获取组件的 HWND
                self.hwnd = self.widget.winfo_id()
                if not self.hwnd:
                    return

                # 防止重复 Hook
                if getattr(self.widget, "_dnd_hooked", False):
                    return
                
                # API Definitions for Hooking
                try:
                    SetWindowLong = ctypes.windll.user32.SetWindowLongPtrW
                    SetWindowLong.argtypes = [ctypes.c_void_p, ctypes.c_int, ctypes.c_void_p]
                    SetWindowLong.restype = ctypes.c_void_p
                except AttributeError:
                    SetWindowLong = ctypes.windll.user32.SetWindowLongW
                    SetWindowLong.argtypes = [ctypes.c_void_p, ctypes.c_int, ctypes.c_long]
                    SetWindowLong.restype = ctypes.c_long

                self.CallWindowProc = ctypes.windll.user32.CallWindowProcW
                self.CallWindowProc.argtypes = [ctypes.c_void_p, ctypes.c_void_p, ctypes.c_uint, ctypes.c_void_p, ctypes.c_void_p]
                self.CallWindowProc.restype = LRESULT
                
                DragAcceptFiles = ctypes.windll.shell32.DragAcceptFiles
                DragAcceptFiles.argtypes = [ctypes.c_void_p, ctypes.c_bool]
                DragAcceptFiles.restype = None
                
                # 尝试解除 UIPI 限制
                hwnd_void = ctypes.c_void_p(self.hwnd)
                try:
                    ChangeWindowMessageFilterEx = ctypes.windll.user32.ChangeWindowMessageFilterEx
                    ChangeWindowMessageFilterEx.argtypes = [ctypes.c_void_p, ctypes.c_uint, ctypes.c_uint, ctypes.c_void_p]
                    ChangeWindowMessageFilterEx.restype = ctypes.c_bool
                    
                    ChangeWindowMessageFilterEx(hwnd_void, 0x0233, 1, None) # WM_DROPFILES
                    ChangeWindowMessageFilterEx(hwnd_void, 0x0049, 1, None) # WM_COPYGLOBALDATA
                except AttributeError:
                    try:
                        ChangeWindowMessageFilter = ctypes.windll.user32.ChangeWindowMessageFilter
                        ChangeWindowMessageFilter.argtypes = [ctypes.c_uint, ctypes.c_uint]
                        ChangeWindowMessageFilter(0x0233, 1)
                        ChangeWindowMessageFilter(0x0049, 1)
                    except:
                        pass

                DragAcceptFiles(hwnd_void, True)
                
                self.new_wnd_proc = WNDPROC(self._wnd_proc)
                GWL_WNDPROC = -4
                
                # 转换函数指针为地址
                new_proc_addr = ctypes.cast(self.new_wnd_proc, ctypes.c_void_p).value
                if not new_proc_addr:
                     new_proc_addr = ctypes.addressof(self.new_wnd_proc)
                     
                self.old_wnd_proc = SetWindowLong(self.hwnd, GWL_WNDPROC, new_proc_addr)
                
                self.widget._dnd_hooked = True
                self.widget._dnd_helper = self
                
            except Exception as e:
                print(f"WindowsDnD hook error: {e}")

        def unhook(self):
            try:
                if self.old_wnd_proc and self.hwnd:
                    try:
                        SetWindowLong = ctypes.windll.user32.SetWindowLongPtrW
                        SetWindowLong.argtypes = [ctypes.c_void_p, ctypes.c_int, ctypes.c_void_p]
                        SetWindowLong.restype = ctypes.c_void_p
                    except AttributeError:
                        SetWindowLong = ctypes.windll.user32.SetWindowLongW
                        SetWindowLong.argtypes = [ctypes.c_void_p, ctypes.c_int, ctypes.c_long]
                        SetWindowLong.restype = ctypes.c_long
                    
                    GWL_WNDPROC = -4
                    SetWindowLong(self.hwnd, GWL_WNDPROC, self.old_wnd_proc)
                    self.old_wnd_proc = None
                    self.widget._dnd_hooked = False
            except Exception as e:
                print(f"WindowsDnD unhook error: {e}")

        def _wnd_proc(self, hwnd, msg, wParam, lParam):
            try:
                if msg == 0x233: # WM_DROPFILES
                    # 仅将数据放入队列，不进行任何 Python 对象操作
                    _dnd_queue.put((self, wParam))
                    return 0
                
                if self.old_wnd_proc:
                    return self.CallWindowProc(self.old_wnd_proc, hwnd, msg, wParam, lParam)
            except Exception:
                pass
            return ctypes.windll.user32.DefWindowProcW(hwnd, msg, wParam, lParam)

        def _handle_drop_internal(self, hDrop):
            """内部处理逻辑，由主循环调用"""
            try:
                # 确保 hDrop 是正确的类型
                if isinstance(hDrop, int):
                    hDrop = ctypes.c_void_p(hDrop)
                
                count = _DragQueryFileW(hDrop, 0xFFFFFFFF, None, 0)
                files = []
                buf = ctypes.create_unicode_buffer(4096)
                for i in range(count):
                    if _DragQueryFileW(hDrop, i, buf, 4096):
                        files.append(buf.value)
                
                pt = wintypes.POINT()
                _DragQueryPoint(hDrop, ctypes.byref(pt))
                ctypes.windll.shell32.DragFinish(hDrop)
                
                if self.on_drop_callback:
                    self.on_drop_callback(files, pt.x, pt.y)
            except Exception as e:
                print(f"DnD Handle error: {e}")

except ImportError:
    # Fallback / Dummy implementation if ctypes not available
    class ClipboardUtils:
        @staticmethod
        def set_files(paths): pass
        @staticmethod
        def get_files(): return []
    WindowsDnD = None

class FileViewerApp(TkinterDnD.Tk if HAS_DND else tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("多项目文件查看器 (级联视图)")
        self.geometry("1400x800")
        
        # Windows DnD 辅助对象列表，用于退出时清理 (尽早初始化)
        self._all_dnd_helpers = []
        
        # 全局视频资源锁：防止多个视频线程同时争抢 OpenCV/FFmpeg 资源导致死锁或崩溃
        self._video_resource_lock = threading.Lock()
        
        # 剪贴板操作状态: None, 'copy', 'cut'
        self._clipboard_op = None 
        self._clipboard_files = [] # 存储文件路径列表

        # 默认路径: 优先当前目录，避免检测网络驱动器导致卡顿
        self.default_start_path = os.getcwd()
        self.base_dir = self.default_start_path
        
        self.history_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "file_viewer_config.json")
        self.history_paths = self._load_history()
        
        if self.history_paths:
            self.base_dir = self.history_paths[0]

        self.columns = [] # 存储 (frame, listbox, current_path) 的列表
        self.preview_frame = None
        self._resize_timer = None # 用于图片缩放防抖
        self._selection_timer = None # 用于列表选择防抖
        self._img_cache = {} # 简单缓存: {path: pil_img}
        self._img_cache_meta = {} # 缓存元数据: {path: {"rotation": 0}}
        self._preload_timer = None # 预加载定时器
        
        # 拖拽状态
        self._is_internal_drag = False
        self._drag_source_col_index = None

        # 快捷键绑定
        self.bind("<Control-o>", lambda e: self._browse_dir())
        self.bind("<F5>", lambda e: self._load_projects())
        self.bind("<Control-q>", lambda e: self.destroy())
        self.bind("<BackSpace>", self._on_backspace) # Backspace 返回上一级
        self.bind("<space>", self._on_space) # Space 播放/暂停
        
        # 全局快捷键
        self.bind_class("Listbox", "<Control-c>", self._on_copy)
        self.bind_class("Listbox", "<Control-v>", self._on_paste)
        self.bind_class("Listbox", "<Delete>", self._on_delete)

        self._init_themes()
        self._apply_theme("默认黑 (Dark)")
        
        # 创建菜单栏
        self._create_menu()
        
        # 自动检查更新 (延迟 3秒，避免影响启动速度)
        if UPDATE_CHECK_URL:
            self.after(3000, lambda: self._check_update(silent=True))

        # self._init_menu() # 替换为侧边工具栏
        self._init_ui()
        
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        
        # 启动 DnD 队列检查循环
        self._check_dnd_queue()
        
        self._load_projects(initial=True)

    def _create_menu(self):
        """创建主菜单"""
        menubar = tk.Menu(self)
        self.config(menu=menubar)
        
        # 1. 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="文件", menu=file_menu)
        file_menu.add_command(label="打开位置... (Ctrl+O)", command=self._browse_dir)
        file_menu.add_command(label="刷新项目 (F5)", command=self._load_projects)
        file_menu.add_separator()
        file_menu.add_command(label="退出 (Ctrl+Q)", command=self.destroy)
        
        # 2. 视图菜单
        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="视图", menu=view_menu)
        
        # 风格子菜单
        theme_menu = tk.Menu(view_menu, tearoff=0)
        view_menu.add_cascade(label="界面风格", menu=theme_menu)
        
        for t_name in self.themes:
            theme_menu.add_command(label=t_name, command=lambda n=t_name: self._apply_theme(n))

        # 3. 设置菜单
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="设置", menu=settings_menu)
        
        # 开机自启
        self.startup_var = tk.BooleanVar(value=self._check_startup())
        settings_menu.add_checkbutton(label="开机自动启动", variable=self.startup_var, command=self._toggle_startup)
        
        # 4. 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="帮助", menu=help_menu)
        help_menu.add_command(label="检查更新", command=lambda: self._check_update(silent=False))
        help_menu.add_separator()
        help_menu.add_command(label="使用说明", command=self._show_usage)
        help_menu.add_separator()
        help_menu.add_command(label="关于", command=self._show_about)

    def _check_update(self, silent=False):
        """检查更新"""
        if not UPDATE_CHECK_URL:
            if not silent:
                messagebox.showinfo("检查更新", "未配置更新服务器地址。\n请在代码中设置 UPDATE_CHECK_URL。")
            return

        def _check_thread():
            try:
                import urllib.request
                # 设置超时
                req = urllib.request.Request(UPDATE_CHECK_URL, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=5) as response:
                    remote_ver = response.read().decode('utf-8').strip()
                
                # 简单的版本比较 (假设是 x.y.z 格式)
                if remote_ver != APP_VERSION:
                    # 发现新版本，切换回主线程显示
                    self.after(0, lambda: self._show_update_dialog(remote_ver))
                else:
                    if not silent:
                        self.after(0, lambda: messagebox.showinfo("检查更新", f"当前已是最新版本 ({APP_VERSION})"))
            except Exception as e:
                if not silent:
                    self.after(0, lambda: messagebox.showerror("检查更新失败", f"无法连接到更新服务器: {e}"))
        
        threading.Thread(target=_check_thread, daemon=True).start()

    def _show_update_dialog(self, remote_ver):
        """显示更新对话框"""
        if messagebox.askyesno("发现新版本", f"发现新版本 v{remote_ver} (当前 v{APP_VERSION})\n\n是否立即下载并更新？"):
            self._perform_update()

    def _perform_update(self):
        """下载并应用更新"""
        if not UPDATE_DOWNLOAD_URL:
             messagebox.showerror("错误", "未配置下载地址 UPDATE_DOWNLOAD_URL")
             return

        pd = ProgressDialog(self, "正在更新")
        
        def _download_thread():
            new_path = ""
            try:
                import urllib.request
                import sys
                
                target_path = os.path.abspath(__file__)
                new_path = target_path + ".new"
                
                pd.update(0, "正在连接服务器...")
                
                # Download with progress
                req = urllib.request.Request(UPDATE_DOWNLOAD_URL, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=15) as response:
                    total_size = int(response.info().get('Content-Length', 0))
                    block_size = 8192
                    downloaded = 0
                    
                    with open(new_path, "wb") as f:
                        while True:
                            if pd.cancelled:
                                raise Exception("用户取消")
                            buffer = response.read(block_size)
                            if not buffer:
                                break
                            f.write(buffer)
                            downloaded += len(buffer)
                            if total_size > 0:
                                percent = (downloaded / total_size) * 100
                                pd.update(percent, f"正在下载: {int(percent)}%")
                
                pd.update(100, "下载完成，正在重启...")
                time.sleep(1)
                
                # Create restart script (Windows batch)
                bat_path = os.path.join(os.path.dirname(target_path), "update_restart.bat")
                with open(bat_path, "w") as f:
                    f.write("@echo off\n")
                    # 等待主进程退出
                    f.write("timeout /t 2 >nul\n") 
                    # 覆盖旧文件
                    f.write(f'move /y "{new_path}" "{target_path}" >nul\n')
                    # 重启应用
                    f.write(f'start "" "{sys.executable}" "{target_path}"\n')
                    # 删除自己
                    f.write(f'(goto) 2>nul & del "{bat_path}"\n')
                
                # Execute and exit
                subprocess.Popen(bat_path, shell=True)
                # 使用 os._exit 确保强制退出，防止 cleanup 逻辑干扰
                os._exit(0)
                
            except Exception as e:
                err_msg = str(e)
                self.after(0, lambda: messagebox.showerror("更新失败", err_msg))
                self.after(0, pd.close)
                # Cleanup
                if new_path and os.path.exists(new_path):
                    try: os.remove(new_path)
                    except: pass

        threading.Thread(target=_download_thread, daemon=True).start()

    def _check_startup(self):
        """检查是否已设置开机自启"""
        try:
            # 使用更安全的异常捕获，防止注册表访问失败导致崩溃
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Run", 0, winreg.KEY_READ)
            try:
                winreg.QueryValueEx(key, "FileViewerPy")
                return True
            except FileNotFoundError:
                return False
            finally:
                winreg.CloseKey(key)
        except Exception as e:
            print(f"Startup check failed: {e}")
            return False

    def _toggle_startup(self):
        """切换开机自启状态"""
        app_name = "FileViewerPy"
        try:
            key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Run", 0, winreg.KEY_ALL_ACCESS)
            
            if self.startup_var.get():
                # 设置开机自启
                if getattr(sys, 'frozen', False):
                    # 如果是打包后的 exe
                    cmd = f'"{sys.executable}"'
                else:
                    # 如果是脚本运行，使用 pythonw.exe 避免黑框 (如果存在)，否则用 python.exe
                    py_exe = sys.executable
                    if "python.exe" in py_exe:
                        pyw = py_exe.replace("python.exe", "pythonw.exe")
                        if os.path.exists(pyw):
                            py_exe = pyw
                    
                    script = os.path.abspath(__file__)
                    cmd = f'"{py_exe}" "{script}"'
                
                winreg.SetValueEx(key, app_name, 0, winreg.REG_SZ, cmd)
                print(f"Startup enabled: {cmd}")
            else:
                # 取消开机自启
                try:
                    winreg.DeleteValue(key, app_name)
                    print("Startup disabled")
                except FileNotFoundError:
                    pass
            
            winreg.CloseKey(key)
        except Exception as e:
            messagebox.showerror("错误", f"无法修改注册表: {e}")
            # 恢复 Checkbox 状态
            self.startup_var.set(not self.startup_var.get())

    def _show_usage(self):
        """显示使用说明"""
        msg = """
【快捷键】
  • 导航: ← / → / ↑ / ↓
  • 进入/打开: Enter (回车)
  • 返回上一级: Backspace (退格)
  • 刷新: F5
  • 打开新位置: Ctrl + O
  • 退出: Ctrl + Q

【文件操作】
  • 复制: Ctrl + C
  • 剪切: Ctrl + X
  • 粘贴: Ctrl + V (支持从外部拖入)
  • 删除: Delete
  • 重命名: 右键菜单 -> 重命名

【功能特色】
  • 米勒列视图 (Miller Columns): 像 macOS Finder 一样层级浏览
  • 预览: 支持文本、代码 (高亮)、图片、Hex 预览
  • 拖拽: 支持文件拖放操作
  • 搜索: 每列顶部支持实时过滤
        """
        messagebox.showinfo("使用说明", msg.strip())

    def _show_about(self):
        """显示关于信息"""
        messagebox.showinfo("关于", f"多项目文件查看器 (File Viewer)\n版本: v{APP_VERSION}\n\n基于 Python Tkinter 构建\n仿 macOS Finder 级联视图体验")

    def _check_dnd_queue(self):
        """定期检查 DnD 队列，处理拖放事件"""
        try:
            while not _dnd_queue.empty():
                try:
                    # 获取 (helper_instance, hDrop)
                    helper, hDrop = _dnd_queue.get_nowait()
                    if helper:
                        helper._handle_drop_internal(hDrop)
                except queue.Empty:
                    break
        except Exception:
            pass
        finally:
            # 每 100ms 检查一次
            self.after(100, self._check_dnd_queue)

    def _on_close(self):
        """退出前清理资源"""
        try:
            for helper in self._all_dnd_helpers:
                try:
                    helper.unhook()
                except:
                    pass
            self._all_dnd_helpers.clear()
        except:
            pass
        self.destroy()

    def _run_command(self, cmd):
        """运行系统命令"""
        try:
            subprocess.Popen(cmd, shell=True)
        except Exception as e:
            messagebox.showerror("错误", f"无法启动程序: {cmd}\n{e}")

    def _open_app(self, app_name):
        """尝试打开常用应用"""
        # === 1. 优先尝试注册表查找 (自动定位) ===
        if app_name == "WeChat":
            # 尝试多个注册表键值
            reg_keys = [
                r"Software\Tencent\WeChat",
                r"Software\Tencent\Weixin"
            ]
            
            for reg_path in reg_keys:
                try:
                    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path)
                    install_dir, _ = winreg.QueryValueEx(key, "InstallPath")
                    winreg.CloseKey(key)
                    
                    if install_dir:
                        # 尝试可能的执行文件名
                        for exe_name in ["WeChat.exe", "Weixin.exe"]:
                            exe_path = os.path.join(install_dir, exe_name)
                            if os.path.exists(exe_path):
                                try:
                                    subprocess.Popen(exe_path)
                                    return
                                except Exception as e:
                                    print(f"Registry launch failed: {e}")
                except Exception:
                    pass

        # === 2. 尝试常用硬编码路径 ===
        # 动态生成路径列表，覆盖更多情况
        wechat_candidates = []
        drives = ["C:", "D:", "E:"]
        bases = [r"\Program Files (x86)", r"\Program Files"]
        # 相对路径模式：(文件夹, 可执行文件)
        patterns = [
            (r"Tencent\WeChat", "WeChat.exe"),
            (r"Tencent\Weixin", "Weixin.exe"),
            (r"Tencent\WeChat", "Weixin.exe"), # 罕见但可能
            (r"Tencent\Weixin", "WeChat.exe")
        ]
        
        for drive in drives:
            for base in bases:
                for folder, exe in patterns:
                    full_path = os.path.join(drive, base, folder, exe)
                    wechat_candidates.append(full_path)

        paths = {
            "WeChat": wechat_candidates,
            "DingTalk": [
                r"C:\Program Files (x86)\DingDing\DingtalkLauncher.exe",
                r"C:\Program Files\DingDing\DingtalkLauncher.exe",
                r"D:\Program Files (x86)\DingDing\DingtalkLauncher.exe",
                r"D:\Program Files\DingDing\DingtalkLauncher.exe"
            ],
            "Quark": [
                r"C:\Program Files\Quark\QuarkBrowser\Quark.exe",
                r"C:\Program Files (x86)\Quark\QuarkBrowser\Quark.exe",
                r"D:\Program Files\Quark\QuarkBrowser\Quark.exe",
                r"C:\Users\Public\Desktop\Quark.lnk" # 尝试快捷方式
            ]
        }
        
        found = False
        if app_name in paths:
            for p in paths[app_name]:
                if os.path.exists(p):
                    try:
                        subprocess.Popen(p)
                        found = True
                        break
                    except Exception as e:
                        print(f"Failed to open {p}: {e}")
        
        if not found:
            # 尝试直接通过名字启动（如果已在 PATH 中）
            try:
                subprocess.Popen(f"start {app_name}", shell=True)
                return
            except:
                pass
            
            # 如果还是找不到，询问用户
            if messagebox.askyesno("提示", f"未找到 {app_name}，是否手动指定其安装位置？"):
                file_path = filedialog.askopenfilename(title=f"选择 {app_name} 可执行文件", filetypes=[("Executable", "*.exe")])
                if file_path:
                    try:
                        subprocess.Popen(file_path)
                    except Exception as e:
                        messagebox.showerror("错误", f"无法启动: {e}")

    def _init_themes(self):
        """初始化主题数据"""
        self.themes = {
            "默认黑 (Dark)": {
                "bg": "#2b2b2b", "fg": "#cccccc", "select_bg": "#4a6984", "select_fg": "#ffffff",
                "darker": "#1e1e1e", "lighter": "#3c3c3c", "border": "#444444", "arrow": "#cccccc",
                "sidebar": "#252526",
                "code": {"keyword": "#cc7832", "string": "#6a8759", "comment": "#808080", "number": "#6897bb", "function": "#ffc66d", "key": "#9876aa"}
            },
            "简约白 (Light)": {
                "bg": "#f0f0f0", "fg": "#333333", "select_bg": "#0078d7", "select_fg": "#ffffff",
                "darker": "#ffffff", "lighter": "#e1e1e1", "border": "#cccccc", "arrow": "#555555",
                "sidebar": "#e6e6e6",
                "code": {"keyword": "#000080", "string": "#008000", "comment": "#808080", "number": "#0000ff", "function": "#795e26", "key": "#a31515"}
            },
            "高级灰 (Gray)": {
                "bg": "#535353", "fg": "#eeeeee", "select_bg": "#404040", "select_fg": "#ffffff",
                "darker": "#333333", "lighter": "#666666", "border": "#222222", "arrow": "#eeeeee",
                "sidebar": "#444444",
                "code": {"keyword": "#ffaa00", "string": "#aaffaa", "comment": "#999999", "number": "#55ffff", "function": "#ffff00", "key": "#ff55ff"}
            },
            "黑客绿 (Matrix)": {
                "bg": "#121212", "fg": "#e0e0e0", "select_bg": "#00695c", "select_fg": "#ffffff",
                "darker": "#002b2b", "lighter": "#1e1e1e", "border": "#333333", "arrow": "#009688",
                "sidebar": "#0a0a0a",
                "code": {"keyword": "#80cbc4", "string": "#a5d6a7", "comment": "#546e7a", "number": "#80deea", "function": "#4db6ac", "key": "#b2dfdb"}
            }
        }

    def _apply_theme(self, theme_name):
        """应用指定主题"""
        if theme_name not in self.themes:
            theme_name = "默认黑 (Dark)"
            
        t = self.themes[theme_name]
        
        self.style = ttk.Style(self)
        self.style.theme_use('clam')
        
        bg = t["bg"]
        fg = t["fg"]
        sel_bg = t["select_bg"]
        sel_fg = t["select_fg"]
        darker = t["darker"]
        lighter = t["lighter"]
        border = t["border"]
        arrow = t["arrow"]
        sidebar_bg = t.get("sidebar", darker)
        
        self.configure(bg=bg)
        
        # 配置通用 ttk 样式
        self.style.configure(".", background=bg, foreground=fg, fieldbackground=darker)
        self.style.configure("TFrame", background=bg)
        self.style.configure("TLabelframe", background=bg, foreground=fg, bordercolor=border)
        self.style.configure("TLabelframe.Label", background=bg, foreground=fg)
        self.style.configure("TLabel", background=bg, foreground=fg)
        
        self.style.configure("TButton", background=lighter, foreground=fg, borderwidth=1, bordercolor=border)
        self.style.map("TButton", background=[("active", darker if theme_name == "简约白 (Light)" else lighter), ("pressed", sel_bg)])
        
        self.style.configure("TCombobox", fieldbackground=darker, background=bg, foreground=fg, arrowcolor=arrow)
        self.style.map("TCombobox", fieldbackground=[("readonly", darker)], selectbackground=[("readonly", sel_bg)], selectforeground=[("readonly", sel_fg)])
        
        # 滚动条
        self.style.configure("TScrollbar", gripcount=0, background=lighter, troughcolor=darker, borderwidth=0, arrowcolor=arrow, arrowsize=18)
        self.style.configure("Vertical.TScrollbar", gripcount=0, background=lighter, troughcolor=darker, borderwidth=0, arrowcolor=arrow, arrowsize=18)
        self.style.configure("Horizontal.TScrollbar", gripcount=0, background=lighter, troughcolor=darker, borderwidth=0, arrowcolor=arrow, arrowsize=18)
        
        self.style.configure("TSeparator", background=border)
        
        # 表格 (Treeview) 样式适配
        self.style.configure("Treeview", 
            background=darker, 
            foreground=fg, 
            fieldbackground=darker,
            borderwidth=0
        )
        self.style.configure("Treeview.Heading", 
            background=lighter, 
            foreground=fg, 
            relief="flat"
        )
        self.style.map("Treeview.Heading",
            background=[("active", darker)]
        )
        
        # 选项卡 (Notebook)
        self.style.configure("TNotebook", background=bg, borderwidth=0)
        tab_padding = [15, 5]
        self.style.configure("TNotebook.Tab", background=lighter, foreground=fg, padding=tab_padding, font=("Segoe UI", 10, "bold"))
        self.style.map("TNotebook.Tab", 
            background=[("selected", sel_bg)], 
            foreground=[("selected", sel_fg)],
            padding=[("selected", tab_padding)]
        )

        # 侧边栏样式
        self.style.configure("Sidebar.TFrame", background=sidebar_bg)
        self.style.configure("Sidebar.TButton", background=sidebar_bg, foreground=fg, borderwidth=0, anchor="w", padding=5)
        self.style.map("Sidebar.TButton", background=[("active", lighter), ("pressed", sel_bg)], foreground=[("pressed", sel_fg)])

        # 语法高亮
        self.code_colors = t["code"]

        # 保存颜色配置
        self.colors = {
            "bg": bg,
            "fg": fg,
            "listbox_bg": darker,
            "listbox_fg": fg,
            "listbox_sel_bg": sel_bg,
            "listbox_sel_fg": sel_fg,
            "text_bg": darker,
            "text_fg": fg,
            "canvas_bg": bg,
            "entry_bg": darker,
            "entry_fg": fg,
            "sidebar_bg": sidebar_bg
        }
        
        # 更新现有非 ttk 组件
        self._update_existing_widgets()

    def _update_existing_widgets(self):
        """刷新已创建的非 ttk 组件颜色"""
        try:
            # 1. Canvas
            if hasattr(self, "canvas"):
                self.canvas.config(bg=self.colors["canvas_bg"])
            
            # 2. 侧边栏 Canvas
            # 需要找到 sidebar 里的 canvas
            # 这里的 sidebar 是在 _init_ui 里创建的，我们没有直接保存 sidebar canvas 的引用
            # 但可以通过遍历 widget 树找到
            
            # 3. Columns (Listboxes, Handles, Text Previews)
            if hasattr(self, "columns"):
                for col in self.columns:
                    # Listbox
                    if col.get("listbox"):
                        col["listbox"].config(
                            bg=self.colors["listbox_bg"],
                            fg=self.colors["listbox_fg"],
                            selectbackground=self.colors["listbox_sel_bg"],
                            selectforeground=self.colors["listbox_sel_fg"]
                        )
                    
                    # Handles (Resize Bars)
                    # Handle 是 container 的子组件，container 是 col["frame"] 的父组件
                    # col["frame"] 是 container 的子组件 (side=left)
                    # handle 是 container 的子组件 (side=right)
                    try:
                        container = col["frame"].master
                        for child in container.winfo_children():
                            if isinstance(child, tk.Frame) and child != col["frame"]:
                                # 这很可能是 handle，因为它是一个普通的 Frame 且不是内容 Frame
                                # 原代码: handle = tk.Frame(container, width=14, bg="#333333", ...)
                                # 只有 handle 是 tk.Frame (内容 frame 是 ttk.LabelFrame)
                                # 或者是 container 本身？container 是 ttk.Frame
                                child.config(bg=self.themes[self.current_theme_name]["lighter"] if hasattr(self, "current_theme_name") else "#333333")
                                # 暂时用 lighter 颜色作为 handle 颜色，或者 border 颜色
                                child.config(bg=self.colors.get("fg", "#555555")) # 稍微显眼一点
                    except:
                        pass
                        
                    # Text Previews (in Preview Column)
                    if col.get("is_preview"):
                         # 找到 Text 组件
                         # 结构: frame -> paned -> preview_area -> Text
                         # 或者 frame -> paned -> info_frame -> Text
                         self._recursive_update_text_bg(col["frame"])

        except Exception as e:
            # 初始化阶段可能某些组件还未创建
            pass

    def _recursive_update_text_bg(self, widget):
        """递归更新 Text 组件背景"""
        for child in widget.winfo_children():
            if isinstance(child, tk.Text):
                child.config(
                    bg=self.colors["text_bg"],
                    fg=self.colors["text_fg"],
                    insertbackground=self.colors["fg"]
                )
                # 还有行号栏 (bg=bg_color, fg=gray)
                # 这是一个特例，之前的代码行号栏背景是硬编码 #2b2b2b
                # 我们尽量更新它
            
            if isinstance(child, tk.Canvas):
                child.config(bg=self.colors["bg"])
                
            self._recursive_update_text_bg(child)

    def _load_history(self) -> list[str]:
        """加载历史记录"""
        if not os.path.exists(self.history_file):
            return [self.default_start_path]
        try:
            with open(self.history_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                paths = data.get("paths", [])
                if not paths:
                    return [self.default_start_path]
                return paths
        except Exception:
            return [self.default_start_path]

    def _save_history(self, new_path: str):
        """保存历史记录"""
        if new_path in self.history_paths:
            self.history_paths.remove(new_path)
        self.history_paths.insert(0, new_path)
        # 只保留最近10条
        self.history_paths = self.history_paths[:10]
        
        try:
            with open(self.history_file, "w", encoding="utf-8") as f:
                json.dump({"paths": self.history_paths}, f, ensure_ascii=False, indent=2)
            # 更新下拉列表
            self.path_combo['values'] = self.history_paths
            self.path_combo.set(new_path)
        except Exception as e:
            print(f"Failed to save config: {e}")

    def _init_ui(self):
        # === 顶部工具架 (Tabbed Tool Rack) ===
        # 使用 Notebook 实现分类显示，节省空间且直观
        tool_notebook = ttk.Notebook(self)
        tool_notebook.pack(fill="x", padx=2, pady=2)
        
        # --- Tab 1: 常用工具 (系统 + 社交) ---
        tab_tools = ttk.Frame(tool_notebook, padding=5)
        tool_notebook.add(tab_tools, text="  🛠️  常用工具  ")
        
        ttk.Button(tab_tools, text="📷 截图", width=8, command=lambda: self._run_command("snippingtool")).pack(side="left", padx=5)
        ttk.Button(tab_tools, text="🧮 计算", width=8, command=lambda: self._run_command("calc")).pack(side="left", padx=5)
        ttk.Separator(tab_tools, orient="vertical").pack(side="left", fill="y", padx=8)
        ttk.Button(tab_tools, text="💬 微信", width=8, command=lambda: self._open_app("WeChat")).pack(side="left", padx=5)
        ttk.Button(tab_tools, text="钉 钉钉", width=8, command=lambda: self._open_app("DingTalk")).pack(side="left", padx=5)

        # --- Tab 2: 浏览器 ---
        tab_browsers = ttk.Frame(tool_notebook, padding=5)
        tool_notebook.add(tab_browsers, text="  🌐  浏览器  ")
        
        ttk.Button(tab_browsers, text="Chrome", width=8, command=lambda: self._run_command("start chrome")).pack(side="left", padx=5)
        ttk.Button(tab_browsers, text="Edge", width=8, command=lambda: self._run_command("start msedge")).pack(side="left", padx=5)
        ttk.Button(tab_browsers, text="夸克", width=8, command=lambda: self._open_app("Quark")).pack(side="left", padx=5)
        ttk.Button(tab_browsers, text="Firefox", width=8, command=lambda: self._run_command("start firefox")).pack(side="left", padx=5)

        # --- Tab 3: AI与娱乐 ---
        tab_webs = ttk.Frame(tool_notebook, padding=5)
        tool_notebook.add(tab_webs, text="  🚀  AI与娱乐  ")
        
        websites = [
            ("即梦AI", "https://jimeng.jianying.com/"),
            ("可灵AI", "https://klingai.kuaishou.com/"),
            ("RunHub", "https://www.runninghub.ai/"),
            ("Liblib", "https://www.liblib.art/"),
            ("豆包", "https://www.doubao.com/"),
            ("B站", "https://www.bilibili.com/"),
            ("百度", "https://www.baidu.com/"),
            ("YouTube", "https://www.youtube.com/")
        ]
        
        import webbrowser
        for name, url in websites:
            ttk.Button(tab_webs, text=name, command=lambda u=url: webbrowser.open(u)).pack(side="left", padx=5)

        # === 顶部地址栏 ===
        top_frame = ttk.Frame(self, padding=5)
        top_frame.pack(fill="x")
        
        ttk.Label(top_frame, text="根目录:").pack(side="left")
        
        # 改为下拉框以显示历史记录
        self.path_combo = ttk.Combobox(top_frame, values=self.history_paths)
        self.path_combo.set(self.base_dir)
        self.path_combo.pack(side="left", fill="x", expand=True, padx=5)
        self.path_combo.bind('<Return>', lambda e: self._load_projects())
        self.path_combo.bind('<<ComboboxSelected>>', lambda e: self._load_projects())
        
        ttk.Button(top_frame, text="浏览...", command=self._browse_dir).pack(side="left")
        ttk.Button(top_frame, text="刷新", command=self._load_projects).pack(side="left", padx=5)

        # === 底部状态栏 ===
        self.global_status_var = tk.StringVar()
        self.global_status_var.set("就绪")
        status_bar = ttk.Label(self, textvariable=self.global_status_var, relief="sunken", anchor="w", padding=(5, 2))
        status_bar.pack(side="bottom", fill="x")

        # === 主体区域 (分割布局: 侧边栏 + Miller Columns) ===
        main_split = tk.PanedWindow(self, orient="horizontal", sashrelief="raised", bg=self.colors["bg"])
        main_split.pack(fill="both", expand=True, padx=5, pady=5)

        # 1. 侧边栏
        sidebar = ttk.Frame(main_split)
        main_split.add(sidebar, stretch="never")
        
        self._init_sidebar(sidebar)

        # 2. Miller Columns 容器
        content_area = ttk.Frame(main_split)
        main_split.add(content_area, stretch="always")
        
        # 水平滚动条
        h_scroll = ttk.Scrollbar(content_area, orient="horizontal")
        h_scroll.pack(side="bottom", fill="x")
        
        # 画布用于承载列
        self.canvas = tk.Canvas(content_area, xscrollcommand=h_scroll.set, bg=self.colors["canvas_bg"], highlightthickness=0)
        self.canvas.pack(side="left", fill="both", expand=True)
        h_scroll.config(command=self.canvas.xview)
        
        self.scroll_frame = ttk.Frame(self.canvas)
        self.scroll_window_id = self.canvas.create_window((0, 0), window=self.scroll_frame, anchor="nw")
        
        # 关键修复：确保 scroll_frame 高度跟随 canvas 高度变化
        def _configure_canvas(event):
            self.canvas.itemconfig(self.scroll_window_id, height=event.height)
            self._auto_fit_preview_column(event.width)
        
        self.canvas.bind("<Configure>", _configure_canvas)
        self.scroll_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # 预览区域 (固定在最右侧，或者作为最后一列，这里我们把预览区域做成一个独立的浮动面板或者固定在最右边？
        # 用户的需求是 "自动再创建一列"，这意味着 Miller Columns (级联列表) 模式。
        # 为了保持预览功能，我们可以在 scroll_frame 的最右侧始终添加一个预览列，或者在点击文件时动态添加。
        # 这里采用：所有列都在 scroll_frame 里动态添加。

    def _init_sidebar(self, parent):
        """初始化侧边栏 (带滚动条)"""
        # 样式
        style = ttk.Style()
        style.configure("Sidebar.TButton", anchor="w", padding=5)
        
        # === 滚动结构 ===
        # 1. 垂直滚动条
        v_scroll = ttk.Scrollbar(parent, orient="vertical")
        v_scroll.pack(side="right", fill="y")
        
        # 2. 画布
        canvas = tk.Canvas(parent, yscrollcommand=v_scroll.set, bg=self.colors["bg"], highlightthickness=0)
        canvas.pack(side="left", fill="both", expand=True)
        v_scroll.config(command=canvas.yview)
        
        # 3. 内部 Frame
        content = ttk.Frame(canvas)
        canvas_window = canvas.create_window((0, 0), window=content, anchor="nw")
        
        # 4. 绑定调整事件
        def _on_frame_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            
        def _on_canvas_configure(event):
            canvas.itemconfig(canvas_window, width=event.width)
            
        content.bind("<Configure>", _on_frame_configure)
        canvas.bind("<Configure>", _on_canvas_configure)
        
        # 5. 滚轮支持
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            
        # 递归绑定滚轮事件到所有子控件
        def _bind_mousewheel(widget):
            widget.bind("<MouseWheel>", _on_mousewheel)
            for child in widget.winfo_children():
                _bind_mousewheel(child)
        
        # --- 内容填充 ---
        lbl = ttk.Label(content, text="📌 收藏夹", font=("Segoe UI", 10, "bold"))
        lbl.pack(fill="x", padx=5, pady=(10, 5))
        
        # 常用路径
        common_paths = [
            ("🖥️ 桌面", os.path.join(os.path.expanduser("~"), "Desktop")),
            ("📄 文档", os.path.join(os.path.expanduser("~"), "Documents")),
            ("⬇️ 下载", os.path.join(os.path.expanduser("~"), "Downloads")),
            ("🎵 音乐", os.path.join(os.path.expanduser("~"), "Music")),
            ("🖼️ 图片", os.path.join(os.path.expanduser("~"), "Pictures")),
            ("🏠 用户主目录", os.path.expanduser("~"))
        ]
        
        for name, path in common_paths:
            if os.path.exists(path):
                btn = ttk.Button(content, text=name, style="Sidebar.TButton", 
                                 command=lambda p=path: self._navigate_to(p))
                btn.pack(fill="x", padx=2, pady=1)
                
        # 磁盘驱动器
        ttk.Separator(content, orient="horizontal").pack(fill="x", padx=5, pady=10)
        ttk.Label(content, text="💾 磁盘", font=("Segoe UI", 10, "bold")).pack(fill="x", padx=5, pady=5)
        
        import string
        drives = []
        if os.name == 'nt':
            drives = [f"{d}:\\" for d in string.ascii_uppercase if os.path.exists(f"{d}:")]
        else:
            drives = ["/"]
            
        for drive in drives:
            btn = ttk.Button(content, text=f"💿 {drive}", style="Sidebar.TButton",
                             command=lambda p=drive: self._navigate_to(p))
            btn.pack(fill="x", padx=2, pady=1)
            
        # 最后应用绑定
        canvas.bind("<MouseWheel>", _on_mousewheel)
        content.bind("<MouseWheel>", _on_mousewheel)
        # 绑定所有子控件
        for child in content.winfo_children():
             _bind_mousewheel(child)

        # 6. 自动调整宽度
        # 等待 idle 任务完成以获取准确的请求大小
        content.update_idletasks()
        req_width = content.winfo_reqwidth()
        # 加上滚动条宽度(约20)和内边距，设置最大最小值限制
        final_width = max(50, min(req_width + 25, 300))
        canvas.config(width=final_width)

    def _navigate_to(self, path):
        """导航到指定路径 (重置视图)"""
        self.path_combo.set(path)
        self._load_projects()

    def _browse_dir(self):
        d = filedialog.askdirectory(initialdir=self.base_dir)
        if d:
            self.path_combo.set(d)
            self._load_projects()

    def _load_projects(self, initial=False):
        path = self.path_combo.get().strip()
        if not path:
            return
            
        if not os.path.exists(path):
            if not initial: 
                messagebox.showerror("错误", "路径不存在")
            return
        
        self.base_dir = path
        self._save_history(path)
        
        # 清除所有列
        self._clear_columns(start_index=0)
        
        # 添加第一列：项目列表
        self._add_column(path, title="项目列表", is_root=True)

    def _clear_columns(self, start_index=0):
        """清除从 start_index 开始的所有列"""
        while len(self.columns) > start_index:
            col = self.columns.pop()
            
            # 停止视频播放
            if "video_stop_event" in col:
                col["video_stop_event"].set()
            
            # 优先检查 Listbox 绑定的 helper
            if "listbox" in col and hasattr(col["listbox"], "_dnd_helper"):
                try:
                    helper = col["listbox"]._dnd_helper
                    helper.unhook()
                    if helper in self._all_dnd_helpers:
                        self._all_dnd_helpers.remove(helper)
                except:
                    pass
            # 兼容旧逻辑
            elif "dnd_helper" in col:
                try:
                    helper = col["dnd_helper"]
                    helper.unhook()
                    if helper in self._all_dnd_helpers:
                        self._all_dnd_helpers.remove(helper)
                except:
                    pass
            
            col["frame"].destroy()

    def _start_resize(self, event, container):
        self._resize_start_x = event.x_root
        self._resize_start_width = container.winfo_width()
        
        # 标记为手动调整过
        if hasattr(container, "is_preview_container"):
             self._preview_manually_resized = True

    def _perform_resize(self, event, container):
        dx = event.x_root - self._resize_start_x
        new_width = max(150, self._resize_start_width + dx)
        container.config(width=new_width)

    def _add_column(self, path, title=None, is_root=False):
        """动态添加一列"""
        # 提前定义 col_index，供后续 lambda 使用，防止 UnboundLocalError
        col_index = len(self.columns)
        
        if title is None:
            title = os.path.basename(path)
            
        # 容器 (包含内容+手柄)
        container = ttk.Frame(self.scroll_frame, width=250)
        container.pack(side="left", fill="y", padx=2, pady=2)
        container.pack_propagate(False)

        # === 修复列宽调整问题 ===
        # 必须先创建并 pack 手柄 (side="right")，然后再 pack 内容 frame (side="left")
        # 这样手柄才不会被内容 frame 的 expand=True 挤出可视区域
        
        # 拖拽手柄 (加宽并增加可见性)
        handle = tk.Frame(container, width=14, bg="#333333", cursor="sb_h_double_arrow")
        handle.pack(side="right", fill="y")
        
        # 内容
        frame = ttk.LabelFrame(container, text=title, padding=2)
        frame.pack(side="left", fill="both", expand=True)
        
        # 手柄视觉反馈
        def _on_handle_enter(e): e.widget.config(bg="#555555")
        def _on_handle_leave(e): e.widget.config(bg="#333333")
        
        handle.bind("<Enter>", _on_handle_enter)
        handle.bind("<Leave>", _on_handle_leave)
        handle.bind("<Button-1>", lambda e: self._start_resize(e, container))
        handle.bind("<B1-Motion>", lambda e: self._perform_resize(e, container))
        
        # 底部状态栏
        status_var = tk.StringVar()
        status_label = ttk.Label(frame, textvariable=status_var, anchor="w", font=("Segoe UI", 8), foreground="#888888")
        status_label.pack(side="bottom", fill="x", padx=2, pady=(2, 0))
        
        # 搜索框
        search_var = tk.StringVar()
        entry = ttk.Entry(frame, textvariable=search_var)
        entry.pack(side="top", fill="x", padx=2, pady=(0, 2))
        
        # 列表容器 (用于包含 Listbox 和 Scrollbar)
        list_container = ttk.Frame(frame)
        list_container.pack(side="top", fill="both", expand=True)
        
        # 使用 Grid 布局以确保滚动条宽度不受列表挤压
        list_container.grid_columnconfigure(0, weight=1)
        list_container.grid_rowconfigure(0, weight=1)

        # 滚动条
        v_scroll = ttk.Scrollbar(list_container, orient="vertical")
        
        def _auto_hide_scroll(first, last):
            first, last = float(first), float(last)
            if first <= 0 and last >= 1:
                v_scroll.grid_remove()
            else:
                v_scroll.grid(row=0, column=1, sticky="ns")
            v_scroll.set(first, last)

        # 列表
        lb = tk.Listbox(list_container, exportselection=False, 
                        bg=self.colors["listbox_bg"], 
                        fg=self.colors["listbox_fg"],
                        selectbackground=self.colors["listbox_sel_bg"],
                        selectforeground=self.colors["listbox_sel_fg"],
                        selectmode=tk.EXTENDED, # 启用多选
                        highlightthickness=0,
                        borderwidth=0,
                        activestyle='none',
                        yscrollcommand=_auto_hide_scroll)
        lb.grid(row=0, column=0, sticky="nsew")

        # 绑定 Space 键以支持视频暂停 (优先于 Listbox 默认选择行为)
        lb.bind("<space>", lambda e: self._on_space_in_listbox(e))

        # === 修复：拦截点击事件，防止点击空白处选中最后一行 ===
        lb.bind("<Button-1>", lambda e: self._check_empty_click(e, lb))
        
        # 启用拖放 (如果是文件夹列)
        
        # 1. Drag Source (拖出) - 始终使用 TkinterDnD
        if HAS_DND:
            lb.drag_source_register(1, DND_FILES)
            lb.dnd_bind('<<DragInitCmd>>', lambda e, idx=col_index: self._on_drag_init(e, idx))
            lb.dnd_bind('<<DragEndCmd>>', lambda e: self._on_drag_end(e))

        # 2. Drop Target (拖入)
        # 策略：如果支持 WindowsDnD，仅使用 WindowsDnD 作为接收端，禁用 TkinterDnD 的接收功能，
        # 避免两个系统同时尝试处理 WM_DROPFILES 导致冲突或失效。
        
        dnd_hook_success = False
        if os.name == 'nt' and WindowsDnD:
            try:
                # 延迟 hook 确保 Listbox HWND 可用
                def _hook_dnd():
                    # 检查列索引和 Listbox 是否仍然有效
                    if col_index >= len(self.columns): return
                    
                    try:
                        # 如果已经 Hook 过，就不再创建新的
                        if getattr(lb, "_dnd_hooked", False):
                            return
                            
                        dnd = WindowsDnD(lb, lambda files, x, y, idx=col_index: self._on_drop(None, idx, files=files))
                        dnd.hook()
                        print(f"WindowsDnD hooked for column {col_index}")
                        
                        # 强引用存储
                        self._all_dnd_helpers.append(dnd)
                    except Exception as e:
                        print(f"Hook failed: {e}")
                        # 如果 Hook 失败，回退到 TkinterDnD
                        if HAS_DND:
                             lb.drop_target_register(DND_FILES)
                             lb.dnd_bind('<<Drop>>', lambda e, idx=col_index: self._on_drop(e, idx))

                self.after(200, _hook_dnd)
                dnd_hook_success = True
            except Exception as e:
                print(f"Failed to setup WindowsDnD: {e}")
        
        # 如果不支持 WindowsDnD 或明确失败（非延迟部分），使用 TkinterDnD
        if not dnd_hook_success and HAS_DND:
            lb.drop_target_register(DND_FILES)
            lb.dnd_bind('<<Drop>>', lambda e, idx=col_index: self._on_drop(e, idx))

        v_scroll.config(command=lb.yview)
        
        # 保存列信息
        # col_index 已在开头定义
        col_data = {
            "frame": container,
            "listbox": lb,
            "path": path,
            "all_items": [], # [(name, full_path, is_dir, size, mtime)]
            "filtered_items": [], 
            "status_var": status_var,
            "search_var": search_var
        }
        self.columns.append(col_data)
        
        # 绑定事件
        lb.bind("<<ListboxSelect>>", lambda e, idx=col_index: self._on_column_select(idx))
        lb.bind("<Double-Button-1>", lambda e, idx=col_index: self._on_column_double_click(idx))
        lb.bind("<Return>", lambda e, idx=col_index: self._on_enter_key(idx))
        lb.bind("<Button-3>", lambda e, idx=col_index: self._on_right_click(e, idx))
        
        # 键盘导航
        lb.bind("<Left>", lambda e, idx=col_index: self._on_key_left(idx))
        lb.bind("<Right>", lambda e, idx=col_index: self._on_key_right(idx))
        
        # 搜索过滤
        def _on_search_change(*args):
            # 防抖处理 (300ms)
            timer_name = f"_search_timer_{col_index}"
            if getattr(self, timer_name, None):
                self.after_cancel(getattr(self, timer_name))
            
            text = search_var.get().lower()
            timer = self.after(300, lambda: self._filter_items(col_index, text))
            setattr(self, timer_name, timer)
            
        search_var.trace("w", _on_search_change)

        # 填充数据 (异步加载)
        self.after(0, lambda: self._reload_column(col_index))


        # 自动滚动到最右边
        # 使用 after_idle 确保在布局更新后执行滚动，解决"无结果"问题且不阻塞 UI
        self.after_idle(lambda: self.canvas.xview_moveto(1.0))

    def _on_column_select(self, col_index):
        """处理列表选中事件 (带防抖)"""
        try:
            col_data = self.columns[col_index]
            sel = col_data["listbox"].curselection()
            
            # 如果没有选中任何项 (例如点击了空白处)，直接返回
            if not sel:
                return

            # 多选逻辑：如果选中多项
            if len(sel) > 1:
                # 1. 更新状态栏
                self.global_status_var.set(f"已选中 {len(sel)} 个项目")
                
                # 2. 清除右侧所有列 (避免歧义)
                if self._selection_timer:
                    self.after_cancel(self._selection_timer)
                    self._selection_timer = None
                self._clear_columns(start_index=col_index + 1)
                return

            # 单选逻辑：继续执行原有逻辑
            # 1. 立即更新状态栏 (轻量操作)
            index = sel[0]
            if index < len(col_data["filtered_items"]):
                _, full_path, _, _, _ = col_data["filtered_items"][index]
                self.global_status_var.set(full_path)
        except:
            pass

        # 2. 防抖处理繁重的 UI 更新 (清除列、加载新内容、预览)
        if self._selection_timer:
            self.after_cancel(self._selection_timer)
            self._selection_timer = None
            
        # 延迟 150ms 执行，给予用户快速浏览的时间
        self._selection_timer = self.after(150, lambda: self._perform_column_select(col_index))

    def _perform_column_select(self, col_index):
        """执行实际的列选择逻辑"""
        # 获取当前列的选择
        try:
            col_data = self.columns[col_index]
            sel = col_data["listbox"].curselection()
            if not sel:
                return
                
            index = sel[0]
            # 注意：这里需要使用 filtered_items 而不是 all_items，因为列表可能被过滤了
            if index >= len(col_data["filtered_items"]):
                return
                
            name, full_path, is_dir, _, _ = col_data["filtered_items"][index]
            
            # 清除该列右侧的所有列
            self._clear_columns(start_index=col_index + 1)
            
            # 更新窗口标题
            self.title(f"{os.path.basename(full_path)} - 多项目文件查看器")
            
            if is_dir:
                # 如果是文件夹，添加新的一列
                self._add_column(full_path)
            else:
                # 如果是文件，显示预览面板
                self._show_preview_column(name, full_path)
                
            # 3. 预加载前后图片的逻辑 (延迟执行，以免阻塞当前显示)
            if not is_dir:
                if self._preload_timer:
                    self.after_cancel(self._preload_timer)
                self._preload_timer = self.after(500, lambda: self._preload_neighbors(col_index, index))
                
        except Exception as e:
            print(f"Selection error: {e}")

    def _preload_neighbors(self, col_index, current_index):
        """预加载当前图片的前后图片"""
        try:
            col_data = self.columns[col_index]
            items = col_data["filtered_items"]
            
            # 预加载前后各 1 张
            to_preload = []
            if current_index > 0:
                to_preload.append(items[current_index - 1])
            if current_index < len(items) - 1:
                to_preload.append(items[current_index + 1])
                
            for _, full_path, is_dir, _, _ in to_preload:
                if not is_dir and is_image_file(full_path) and full_path not in self._img_cache:
                    if HAS_PIL:
                        # 启动后台线程预加载
                        import threading
                        threading.Thread(target=self._load_image_worker, args=(full_path,), daemon=True).start()
        except:
            pass

    def _on_column_double_click(self, col_index):
        col_data = self.columns[col_index]
        sel = col_data["listbox"].curselection()
        if not sel: return
        
        index = sel[0]
        # 使用 filtered_items
        name, full_path, is_dir, _, _ = col_data["filtered_items"][index]
        
        if not is_dir:
            # 文件：系统打开
            self._open_system_file_path(full_path)

    def _auto_fit_preview_column(self, canvas_width):
        """自动调整预览列宽度以填满剩余空间"""
        # 如果用户手动调整过，就不再自动调整
        if getattr(self, "_preview_manually_resized", False):
            return

        if not self.columns:
            return
            
        last_col = self.columns[-1]
        # 只有当最后一列是预览列时才调整
        if not last_col.get("is_preview", False):
            return
            
        # 计算前面所有列占用的宽度
        used_width = 0
        for i in range(len(self.columns) - 1):
            col = self.columns[i]
            if col["frame"].winfo_exists():
                used_width += col["frame"].winfo_width() + 4 # padx=2 * 2
        
        # 最后一列本身的 padding
        used_width += 4
        
        # 计算目标宽度
        target_width = max(225, canvas_width - used_width)
        
        # 如果当前宽度与目标宽度差异较大，则更新
        # 注意：这里需要避免频繁微小更新导致的抖动
        if last_col["frame"].winfo_exists():
            current_width = last_col["frame"].winfo_width()
            # 只有在初始化(width=1)或者差异较大时才更新
            if current_width <= 1 or abs(target_width - current_width) > 5:
                last_col["frame"].config(width=int(target_width))

    def _load_image_worker(self, path, callback=None, label=None, preview_area=None, initial_width=None):
        """实际执行图片加载的 Worker 函数"""
        try:
            # 检查缓存 (双重检查)
            if path in self._img_cache:
                pil_img = self._img_cache[path]
            else:
                pil_img = Image.open(path)
                # 预处理：保留高分辨率 (2560px)，满足用户对清晰度的要求
                # 既保留细节又防止 8K/RAW 图片撑爆内存
                pil_img.thumbnail((2560, 2560)) 
                
                # 放入缓存
                if len(self._img_cache) > 30: # 增加缓存容量
                    self._img_cache.pop(next(iter(self._img_cache)))
                self._img_cache[path] = pil_img
            
            # 如果有回调，则在主线程执行
            if callback:
                self.after(0, lambda: callback(pil_img, label, preview_area, initial_width))
        except Exception as e:
            if label:
                self.after(0, lambda: label.config(text=f"图片加载失败: {e}"))

    def _load_image_threaded(self, path, label, preview_area, initial_width):
        """后台线程加载图片"""
        import threading
        threading.Thread(target=self._load_image_worker, 
                         args=(path, self._update_image_preview, label, preview_area, initial_width), 
                         daemon=True).start()

    def _update_image_preview(self, pil_img, label, preview_area, initial_width):
        """图片加载完成后的回调"""
        if not label.winfo_exists():
            return
            
        # 绑定 Resize
        preview_area.bind("<Configure>", lambda e, l=label, img=pil_img: self._resize_preview_image(e, l, img))
        
        # 初始显示
        est_w = initial_width - 20
        est_h = self.winfo_height() // 2 
        if est_w > 50:
            self._perform_image_resize(est_w, est_h, label, pil_img, resample=Image.Resampling.NEAREST)

    def _resize_preview_image(self, event, label, original_pil_img):
        """根据容器大小动态调整图片 (防抖 + 防死循环)"""
        if not original_pil_img:
            return

        # 获取当前尺寸
        target_w = event.width - 20
        target_h = event.height - 20
        
        if target_w < 50 or target_h < 50:
            return

        # 检查尺寸是否真的发生了显著变化 (防止 Configure 死循环导致界面卡死)
        last_size = getattr(label, "_last_resize_size", (0, 0))
        if abs(target_w - last_size[0]) < 5 and abs(target_h - last_size[1]) < 5:
            return
            
        label._last_resize_size = (target_w, target_h)

        # 1. 取消之前的定时器
        if self._resize_timer:
            self.after_cancel(self._resize_timer)
            self._resize_timer = None

        # 3. 设置新的定时器 (50ms)
        self._resize_timer = self.after(50, lambda: self._perform_image_resize(target_w, target_h, label, original_pil_img))

    def _perform_image_resize(self, target_w, target_h, label, original_pil_img, resample=Image.Resampling.BILINEAR):
        """异步执行实际的图片缩放"""
        def _resize_task():
            try:
                if not label.winfo_exists(): return
                
                w, h = original_pil_img.size
                
                # 计算缩放比例
                ratio_w = target_w / w
                ratio_h = target_h / h
                ratio = min(ratio_w, ratio_h)
                
                if ratio < 1.0:
                    new_w = int(w * ratio)
                    new_h = int(h * ratio)
                else:
                    new_w = w
                    new_h = h
                    
                # 缩放
                resized_pil = original_pil_img.resize((new_w, new_h), resample)
                
                # 回到主线程更新 UI
                self.after(0, lambda: self._update_resized_image(label, resized_pil))
            except Exception as e:
                print(f"Resize error: {e}")
        
        threading.Thread(target=_resize_task, daemon=True).start()

    def _update_resized_image(self, label, pil_img):
        """主线程更新图片"""
        try:
            if not label.winfo_exists(): return
            tk_img = ImageTk.PhotoImage(pil_img)
            label.config(image=tk_img)
            label.image = tk_img
        except Exception as e:
            print(f"Update image error: {e}")

    def _video_thread(self, path, queue, stop_event, control_state):
        """后台线程：读取视频帧并放入队列"""
        # 使用全局锁保护 OpenCV 初始化，防止旧线程还没释放完，新线程就抢占
        if not hasattr(self, "_video_resource_lock"):
             self._video_resource_lock = threading.Lock()
             
        # 在线程内部导入 cv2，确保安全
        try:
            import cv2
            import gc
        except ImportError:
            return

        # 尝试导入音频播放器 (ffpyplayer)
        player = None
        try:
            from ffpyplayer.player import MediaPlayer
            # vn=True: 禁用视频解码，只播放音频
            player = MediaPlayer(path, ff_opts={'vn': True})
        except:
            pass

        try:
            import math
            
            # 关键：加锁初始化 VideoCapture
            with self._video_resource_lock:
                if stop_event.is_set(): return
                cap = cv2.VideoCapture(path)

            # 1. 检查视频是否成功打开
            if not cap.isOpened():
                queue.put(("ERROR", "无法打开视频文件 (OpenCV open failed)"))
                return

            fps = cap.get(cv2.CAP_PROP_FPS)
            if fps <= 0 or fps > 120: fps = 25
            control_state['fps'] = fps
            delay = 1.0 / fps
            
            # 获取视频信息
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            if width == 0 or height == 0:
                 queue.put(("ERROR", "视频尺寸无效 (0x0)"))
                 return

            gcd_val = math.gcd(width, height)
            ar_w = width // gcd_val
            ar_h = height // gcd_val
            
            control_state['info'] = {
                'resolution': f"{width}x{height}",
                'aspect_ratio': f"{ar_w}:{ar_h}",
                'fps': fps
            }
            
            # 获取总时长
            total_frames = cap.get(cv2.CAP_PROP_FRAME_COUNT)
            duration = total_frames / fps if fps > 0 else 0
            control_state['duration'] = duration
            
            # 计算跳帧步长 (基于分辨率和FPS动态调整)
            # 策略: 分辨率越高，目标FPS越低，以减轻解码压力
            total_pixels = width * height
            
            if total_pixels > 3840 * 2160: # > 4K
                target_fps = 20.0 # 4K+ 视频限制为 20FPS
            elif total_pixels > 1920 * 1080: # > 1080p (2K/4K)
                target_fps = 24.0 # 2K/4K 视频限制为 24FPS
            else:
                target_fps = 30.0 # 普通视频 30FPS
            
            skip_step = 1
            if fps > target_fps:
                skip_step = int(fps / target_fps)
                if skip_step < 1: skip_step = 1
                
                # 更新控制 FPS 为实际播放 FPS
                real_fps = fps / skip_step
                control_state['fps'] = real_fps
                delay = 1.0 / real_fps
            
            frame_counter = 0

            while not stop_event.is_set():
                # 1. 处理跳转请求
                did_seek = False
                if 'seek_req' in control_state:
                    seek_pos = control_state.pop('seek_req')
                    try:
                        cap.set(cv2.CAP_PROP_POS_MSEC, seek_pos * 1000)
                        did_seek = True
                        if player: player.seek(seek_pos, relative=False)
                    except:
                        pass
                
                # 2. 处理暂停 (如果刚跳转过，必须读取一帧以更新画面，即使是暂停状态)
                is_paused = control_state.get('paused', False)
                speed = control_state.get('speed', 1.0)
                
                if player: 
                    # 如果倍速不是 1.0，暂停音频以避免严重不同步 (简单策略)
                    if speed != 1.0:
                        player.set_pause(True)
                    else:
                        player.set_pause(is_paused)

                if is_paused and not did_seek:
                    time.sleep(0.05)
                    continue

                # 音量控制
                if player:
                    target_vol = 0.0 if control_state.get("muted", False) else control_state.get("volume", 1.0)
                    player.set_volume(target_vol)

                start_time = time.time()
                
                # === 跳帧处理 ===
                # 如果不需要精确的每一帧 (预览模式)，跳过中间帧以降低解码负载
                try:
                    if skip_step > 1:
                        for _ in range(skip_step - 1):
                            cap.grab() # 快速抓取但不解码
                            frame_counter += 1

                    ret, frame = cap.read()
                    frame_counter += 1
                except Exception as e:
                    print(f"Read frame error: {e}")
                    ret = False
                
                if not ret:
                    # 播放结束处理
                    if stop_event.is_set(): break
                    
                    # 如果是刚开始就读取失败，可能是文件损坏或格式不支持
                    if cap.get(cv2.CAP_PROP_POS_FRAMES) == 0:
                        queue.put(("ERROR", "无法读取视频帧 (可能是格式不支持)"))
                        break
                        
                    # 循环播放
                    cap.set(cv2.CAP_PROP_POS_FRAMES, 0)
                    if player:
                        player.seek(0, relative=False)
                    time.sleep(0.01)
                    continue

                # === 音画同步：丢帧判断 (Drop Frames) ===
                # 仅在非暂停、非拖拽且有音频且倍速为1.0时检查
                if player and not control_state.get('paused') and not control_state.get('dragging') and control_state.get('speed', 1.0) == 1.0:
                    audio_pts = player.get_pts()
                    if audio_pts is not None:
                        video_pts = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0
                        # 如果视频落后音频超过 0.15秒，丢弃此帧以追赶
                        if video_pts < audio_pts - 0.15:
                            continue
                    
                # 缩放处理
                h, w = frame.shape[:2]
                
                # 动态获取目标尺寸 (来自 UI 线程更新)
                target_w = control_state.get('view_width', 600)
                target_h = control_state.get('view_height', 400)
                
                # === 画质控制逻辑 ===
                # 根据选择的画质模式，强制限制最大分辨率
                # 这比简单的 view_width 更有效，因为它能减少像素处理量
                quality_mode = control_state.get("quality_mode", "Auto")
                max_dim = 0 # 0 表示不限制 (Auto)
                
                if quality_mode == "360P":
                    max_dim = 480 # 限制长边为 480px (通常 360p 是 480x360)
                elif quality_mode == "720P":
                    max_dim = 1280
                elif quality_mode == "1080P":
                    max_dim = 1920
                elif quality_mode == "4K":
                    max_dim = 3840
                
                if max_dim > 0:
                     # 如果当前目标尺寸超过了限制，强制缩小
                     # 但我们也要考虑容器本身就很小的情况，所以取最小值
                     target_w = min(target_w, max_dim)
                     target_h = min(target_h, max_dim)

                # 避免过小
                target_w = max(200, target_w)
                target_h = max(150, target_h)
                
                # Aspect Fit: 保持比例缩放以适应容器
                if w > 0 and h > 0:
                    scale_w = target_w / w
                    scale_h = target_h / h
                    scale = min(scale_w, scale_h) # 取较小的比例，确保完全放入
                    
                    new_w = int(w * scale)
                    new_h = int(h * scale)
                    
                    try:
                        # 优化：使用 INTER_NEAREST 提升性能 (解决卡顿) - 预览不需要高质量插值
                        # 如果画质要求高 (4K/1080P)，可以使用线性插值以获得更好效果
                        interp = cv2.INTER_NEAREST
                        if quality_mode in ["1080P", "4K"]:
                            interp = cv2.INTER_LINEAR
                            
                        frame = cv2.resize(frame, (new_w, new_h), interpolation=interp)
                    except:
                        pass
                    
                # BGR -> RGB
                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(rgb)
                
                # 获取当前时间
                current_pos = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0
                current_frame_idx = int(cap.get(cv2.CAP_PROP_POS_FRAMES))
                
                # 放入队列 (图片, 当前时间, 当前帧索引)
                try:
                    # 如果队列满了，强制移除最旧的一帧，确保最新帧能放入，且不阻塞线程
                    if queue.full():
                        try: queue.get_nowait()
                        except: pass
                        
                    queue.put_nowait((img, current_pos, current_frame_idx))
                except:
                    pass
                    
                # 3. 帧率与同步控制
                if control_state.get('dragging', False) or control_state.get('paused', False):
                    time.sleep(0.05)
                else:
                    speed = control_state.get("speed", 1.0)
                    if speed <= 0: speed = 1.0
                    effective_delay = delay / speed

                    # 音频同步等待 (仅在1.0倍速且有音频时)
                    # 修复：禁用强制音画同步等待，因为它会导致 1.0x 倍速下严重卡顿
                    # 原因：当 diff 较小时跳过了 FPS 等待，导致播放速度不均匀；
                    # 且 get_pts() 可能存在抖动。现在统一使用 standard FPS wait 保证流畅度。
                    synced_audio = False
                    # if player and speed == 1.0:
                    #     audio_pts = player.get_pts()
                    #     if audio_pts is not None:
                    #         video_pts = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0
                    #         diff = video_pts - audio_pts
                    #         
                    #         # 视频超前音频（画面快了），需要等待音频追上来
                    #         # 优化：放宽同步阈值 (0.01s -> 0.05s) 以避免微小差异导致的频繁卡顿
                    #         if diff > 0.05:
                    #             wait_t = min(diff, 0.5)
                    #             time.sleep(wait_t)
                    #         
                    #         synced_audio = True
                    
                    # 如果没有音频同步，使用标准 FPS 等待
                    if not synced_audio:
                        elapsed = time.time() - start_time
                        wait = effective_delay - elapsed
                        if wait > 0:
                            time.sleep(wait)
            
            # 关键：释放资源
            # 先关闭音频 (ffpyplayer)，移出锁外以防阻塞
            if player:
                try: player.close_player()
                except: pass
            
            # 再关闭视频 (OpenCV)，需加锁保护
            with self._video_resource_lock:
                cap.release()
                
                # 显式清理
                del cap
                del player
                gc.collect()
            
        except Exception as e:
            print(f"Video thread error: {e}")

    def _update_video_label(self, label, queue, stop_event, control_state, slider, time_label, play_btn, info_label=None):
        """主线程更新视频帧和进度条"""
        if stop_event.is_set():
            return
            
        if not label.winfo_exists():
            stop_event.set()
            return
        
        try:
            # 1. 尝试获取数据
            try:
                item = queue.get_nowait()
            except:
                # 队列为空，极短时间重试 (5ms) 以保证流畅度
                self.after(5, lambda: self._update_video_label(label, queue, stop_event, control_state, slider, time_label, play_btn, info_label))
                return

            # 优化：不再激进地消费掉所有积压的帧，而是按顺序播放
            # 生产者已有 frame skipping 逻辑，消费者应尽量平滑展示每一帧
            
            # 3. 错误处理
            if isinstance(item, tuple) and item[0] == "ERROR":
                label.config(text=f"视频错误: {item[1]}", foreground="red")
                return # 停止更新

            # 4. 解析数据
            # 兼容旧版本元组长度 (img, pos) 或新版本 (img, pos, frame_idx)
            if len(item) == 3:
                img, current_pos, current_frame = item
            else:
                img, current_pos = item
                current_frame = 0
            
            # 5. 更新画面
            try:
                tk_img = ImageTk.PhotoImage(img)
                
                # 保存当前 PIL 图片供截图使用
                control_state["current_image"] = img
                
                # 判断是否在全屏模式
                if "fs_label" in control_state:
                    try:
                        if control_state["fs_label"].winfo_exists():
                            control_state["fs_label"].configure(image=tk_img)
                            control_state["fs_label"].image = tk_img
                    except:
                        # 如果全屏窗口意外关闭
                        if "fs_label" in control_state: del control_state["fs_label"]
                        label.configure(image=tk_img)
                        label.image = tk_img
                else:
                    if label.winfo_exists():
                        label.configure(image=tk_img)
                        label.image = tk_img 
            except Exception as e:
                print(f"Update image failed: {e}")
            
            # 更新状态 (如果用户没有正在拖拽)
            if not control_state.get('dragging', False):
                control_state['current'] = current_pos
                slider.set(current_pos)
                
            # 更新时间标签 (精确到帧)
            duration = control_state.get('duration', 0)
            fps = control_state.get('fps', 25)
            
            cur_str = self._format_time_frames(current_pos, fps)
            tot_str = self._format_time_frames(duration, fps)
            time_label.config(text=f"{cur_str} / {tot_str}")
            
            # 更新视频信息 (仅需设置一次)
            if info_label and 'info' in control_state and not getattr(info_label, 'info_set', False):
                info = control_state['info']
                info_text = f"分辨率: {info['resolution']} | 比例: {info['aspect_ratio']} | 帧率: {info['fps']:.2f} FPS"
                info_label.config(text=info_text)
                info_label.info_set = True
            
            # 更新按钮文本 (可选，防止状态不一致)
            play_text = "▶" if control_state.get('paused') else "⏸"
            if play_btn.cget("text") != play_text:
                play_btn.config(text=play_text)
            
            # 标记任务完成
            try: queue.task_done()
            except: pass

        except Exception as e:
            print(f"Update video error: {e}")
            pass
            
        # 提高刷新率 (15ms -> ~60FPS) 以获得更流畅的视觉体验
        self.after(15, lambda: self._update_video_label(label, queue, stop_event, control_state, slider, time_label, play_btn, info_label))

    def _format_time(self, seconds):
        """格式化时间 MM:SS"""
        m = int(seconds // 60)
        s = int(seconds % 60)
        return f"{m:02d}:{s:02d}"

    def _format_time_frames(self, seconds, fps):
        """格式化时间 MM:SS:FF"""
        if fps <= 0: fps = 25
        m = int(seconds // 60)
        s = int(seconds % 60)
        f = int((seconds - int(seconds)) * fps)
        return f"{m:02d}:{s:02d}:{f:02d}"

    def _highlight_syntax(self, text_widget, ext):
        """简单的语法高亮 (基于正则)"""
        try:
            import re
            
            # 1. 配置标签颜色
            for tag, color in self.code_colors.items():
                text_widget.tag_configure(tag, foreground=color)
            
            # 获取内容
            content = text_widget.get("1.0", "end")
            
            # 2. 定义规则 (正则, 标签)
            rules = []
            
            if ext in ['.py', '.pyw']:
                rules = [
                    (r'\b(def|class|return|if|else|elif|while|for|in|import|from|try|except|with|as|pass|break|continue|lambda|await|async)\b', 'keyword'),
                    (r'#.*', 'comment'),
                    (r'""".*?"""', 'string'), # 简单多行字符串
                    (r"'''.*?'''", 'string'),
                    (r'(".*?"|\'.*?\')', 'string'),
                    (r'\b\d+\b', 'number'),
                    (r'\bdef\s+(\w+)', 'function'),
                    (r'\bclass\s+(\w+)', 'function')
                ]
            elif ext in ['.json']:
                rules = [
                    (r'(".*?")\s*:', 'key'), # JSON Key
                    (r':\s*(".*?")', 'string'), # JSON String Value
                    (r'\b(true|false|null)\b', 'keyword'),
                    (r'\b\d+\b', 'number')
                ]
            elif ext in ['.js', '.ts', '.jsx', '.tsx', '.html', '.css']:
                rules = [
                    (r'\b(function|var|let|const|return|if|else|for|while|import|export|default|class|this|new)\b', 'keyword'),
                    (r'//.*', 'comment'),
                    (r'/\*.*?\*/', 'comment'),
                    (r'(".*?"|\'.*?\'|`.*?`)', 'string'),
                    (r'\b\d+\b', 'number')
                ]
            elif ext in ['.md', '.markdown']:
                rules = [
                    (r'^#+ .*', 'function'), # Headers
                    (r'\*\*.*?\*\*', 'keyword'), # Bold
                    (r'\*.*?\*', 'string'), # Italic
                    (r'`.*?`', 'number'), # Inline code
                    (r'\[.*?\]\(.*?\)', 'key'), # Links
                    (r'^\s*[-*] ', 'comment'), # List bullets
                    (r'> .*', 'comment'), # Blockquote
                    (r'^```.*', 'string') # Code block fence
                ]
            elif ext in ['.yaml', '.yml']:
                rules = [
                    (r'^[a-zA-Z0-9_-]+:', 'key'), # Key
                    (r':\s*.*', 'string'), # Value
                    (r'^\s*-\s+', 'keyword'), # List item
                    (r'#.*', 'comment')
                ]
            elif ext in ['.ini', '.toml', '.cfg']:
                rules = [
                    (r'^\[.*\]', 'function'), # Section
                    (r'^[a-zA-Z0-9_-]+', 'key'), # Key
                    (r'=', 'keyword'), 
                    (r'#.*', 'comment'),
                    (r';.*', 'comment')
                ]
            elif ext in ['.md', '.markdown']:
                rules = [
                    (r'^#+ .*', 'function'), # Headers
                    (r'\*\*.*?\*\*', 'keyword'), # Bold
                    (r'\*.*?\*', 'string'), # Italic
                    (r'`.*?`', 'number'), # Inline code
                    (r'\[.*?\]\(.*?\)', 'key'), # Links
                    (r'^\s*[-*] ', 'comment'), # List bullets
                    (r'> .*', 'comment'), # Blockquote
                    (r'```[\s\S]*?```', 'string') # Code block
                ]
            elif ext in ['.yaml', '.yml']:
                rules = [
                    (r'^[a-zA-Z0-9_-]+:', 'key'), # Key
                    (r':\s*.*', 'string'), # Value
                    (r'^\s*-\s+', 'keyword'), # List item
                    (r'#.*', 'comment')
                ]
            elif ext in ['.ini', '.toml', '.cfg']:
                rules = [
                    (r'^\[.*\]', 'function'), # Section
                    (r'^[a-zA-Z0-9_-]+', 'key'), # Key
                    (r'=', 'keyword'), 
                    (r'#.*', 'comment'),
                    (r';.*', 'comment')
                ]
            
            # 3. 应用规则 (分块处理以避免卡顿)
            if not rules:
                return

            def _apply_rules():
                for pattern, tag in rules:
                    for match in re.finditer(pattern, content):
                        # 计算 Tkinter 索引
                        # 注意：re.finditer 返回的是绝对偏移量，需要转换为 line.col 格式
                        # 这里为了性能，简化处理：仅当文件较小时使用精确匹配，
                        # 或者使用 search 逐行扫描？
                        # 逐行扫描在 Tkinter 中更高效
                        pass
                
                # 重新实现：逐行扫描比全文正则更适合 Text 组件
                count = tk.IntVar()
                for pattern, tag in rules:
                    # 使用 Text 组件内置的 search 功能 (支持正则)
                    start = "1.0"
                    while True:
                        pos = text_widget.search(pattern, start, stopindex="end", count=count, regexp=True)
                        if not pos:
                            break
                        end = f"{pos} + {count.get()}c"
                        text_widget.tag_add(tag, pos, end)
                        start = end
            
            # 延迟执行高亮，优先显示文本
            self.after(100, _apply_rules)
            
        except Exception as e:
            print(f"Highlight error: {e}")

    def _show_image_preview(self, container, path, initial_width):
        """显示图片预览 (支持缩放拖拽)"""
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=0) # 工具栏
        container.grid_rowconfigure(1, weight=1) # 画布
        
        # 使用主题背景色
        bg_color = self.colors.get("bg", "#1e1e1e")
        fg_color = self.colors.get("fg", "#ffffff")
        
        # === 1. 工具栏 ===
        toolbar = ttk.Frame(container, padding=2)
        toolbar.grid(row=0, column=0, sticky="ew")
        
        # 按钮 (回调在 _init_canvas 中绑定)
        btn_rotate_left = ttk.Button(toolbar, text="↺ 向左旋转", width=10, state="disabled")
        btn_rotate_left.pack(side="left", padx=2)
        
        btn_rotate_right = ttk.Button(toolbar, text="↻ 向右旋转", width=10, state="disabled")
        btn_rotate_right.pack(side="left", padx=2)
        
        # === 2. 画布 ===
        canvas = tk.Canvas(container, bg=bg_color, highlightthickness=0)
        canvas.grid(row=1, column=0, sticky="nsew")
        
        # 提示文本
        canvas.create_text(initial_width//2, 200, text="Loading...", fill=fg_color, tags="loading", anchor="center")
        
        if not HAS_PIL:
             canvas.itemconfigure("loading", text="需要安装 Pillow (PIL) 库以支持图片缩放预览")
             return

        def _load_and_show():
            try:
                # 1. Load Image
                pil_img = Image.open(path)
                
                # Cache (Simple)
                if not hasattr(self, "_img_cache"): self._img_cache = {}
                self._img_cache[path] = pil_img
                
                # Apply Rotation from Meta
                if not hasattr(self, "_img_cache_meta"): self._img_cache_meta = {}
                if path in self._img_cache_meta and "rotation" in self._img_cache_meta[path]:
                    angle = self._img_cache_meta[path]["rotation"]
                    if angle != 0:
                        pil_img = pil_img.rotate(angle, expand=True)
                
                # 2. Initial Scale
                w, h = pil_img.size
                cw = max(initial_width, 100)
                # 尝试获取真实高度，如果获取不到则给一个更大的默认值
                ch = container.winfo_height()
                if ch <= 1: ch = 600 
                
                scale = min(cw/w, ch/h) if w > 0 and h > 0 else 1.0
                scale = min(scale, 1.0) # Don't upscale initially
                
                self.after(0, lambda: _init_canvas(pil_img, scale))
            except Exception as e:
                def show_error(msg):
                    canvas.delete("loading")
                    canvas.create_text(initial_width//2, 200, text=msg, fill="red", width=initial_width-20, anchor="center")
                self.after(0, lambda: show_error(f"加载失败: {e}"))
        
        def _init_canvas(pil_img, initial_scale):
            canvas.delete("loading")
            
            # 启用按钮
            btn_rotate_left.config(state="normal")
            btn_rotate_right.config(state="normal")
            
            # State
            state = {
                "img": pil_img, 
                "scale": initial_scale, 
                "x": 0, "y": 0,
                "tk_img": None,
                "last_cw": 0,
                "last_ch": 0,
                "manual_zoom": False # 是否手动缩放过
            }
            
            # Rotation Logic
            def rotate_image(direction):
                # direction: 1 (left/ccw +90), -1 (right/cw -90)
                # PIL rotate: positive is CCW
                angle = 90 if direction == 1 else -90
                
                # Update Meta
                if not hasattr(self, "_img_cache_meta"): self._img_cache_meta = {}
                if path not in self._img_cache_meta: self._img_cache_meta[path] = {"rotation": 0}
                
                current_rot = self._img_cache_meta[path].get("rotation", 0)
                new_rot = (current_rot + angle) % 360
                self._img_cache_meta[path]["rotation"] = new_rot
                
                # Rotate current image
                # expand=True ensures the new image fits the rotated bounding box
                rotated_img = state["img"].rotate(angle, expand=True)
                state["img"] = rotated_img
                state["tk_img"] = None # Force redraw
                
                draw(force_resize=True)
            
            # Bind buttons
            btn_rotate_left.configure(command=lambda: rotate_image(1))
            btn_rotate_right.configure(command=lambda: rotate_image(-1))
            
            last_pos = [0, 0]
            
            def draw(force_resize=True):
                cw = canvas.winfo_width()
                ch = canvas.winfo_height()
                
                # 如果 canvas 尚未显示，延迟或使用 container 大小
                if cw <= 1 or ch <= 1:
                    cw = container.winfo_width()
                    ch = container.winfo_height()
                    # 如果还是太小，可能还没 pack 好，但我们先尝试绘制，依赖 <Configure> 修正
                    if cw <= 1: cw = initial_width
                    if ch <= 1: ch = 600
                
                # Resize Image
                w, h = state["img"].size
                nw = int(w * state["scale"])
                nh = int(h * state["scale"])
                if nw <= 0 or nh <= 0: return
                
                # 只有当缩放改变时才重新采样 (优化性能)
                if force_resize or state["tk_img"] is None:
                    try:
                        # 使用 BILINEAR 替代 LANCZOS 以提升性能 (解决卡顿)
                        img_resized = state["img"].resize((nw, nh), Image.Resampling.BILINEAR)
                        tk_img = ImageTk.PhotoImage(img_resized)
                        state["tk_img"] = tk_img
                        # 绑定到 canvas 防止 GC
                        canvas.image = tk_img 
                    except Exception as e:
                        print(f"Resize error: {e}")
                        return
                
                tk_img = state["tk_img"]
                
                canvas.delete("img")
                
                # Center calculation
                cx = cw // 2 + state["x"]
                cy = ch // 2 + state["y"]
                
                canvas.create_image(cx, cy, image=tk_img, anchor="center", tags="img")
                
                state["last_cw"] = cw
                state["last_ch"] = ch
                
            draw()
            
            # Events
            def on_wheel(e):
                state["manual_zoom"] = True
                factor = 1.1 if e.delta > 0 else 0.9
                state["scale"] *= factor
                draw(force_resize=True)
                
            def start_pan(e):
                last_pos[0] = e.x
                last_pos[1] = e.y
            
            def pan(e):
                dx = e.x - last_pos[0]
                dy = e.y - last_pos[1]
                state["x"] += dx
                state["y"] += dy
                # 仅移动，不重绘
                canvas.move("img", dx, dy)
                last_pos[0] = e.x
                last_pos[1] = e.y
            
            def reset(e):
                state["manual_zoom"] = False
                # Reset to fit current window
                cw = canvas.winfo_width()
                ch = canvas.winfo_height()
                w, h = state["img"].size
                if w > 0 and h > 0:
                    new_scale = min(cw/w, ch/h)
                    state["scale"] = min(new_scale, 1.0)
                else:
                    state["scale"] = initial_scale
                
                state["x"] = 0
                state["y"] = 0
                draw(force_resize=True)
            
            def on_configure(e):
                cw, ch = e.width, e.height
                # 只有当大小真正改变时才重绘
                if abs(cw - state["last_cw"]) > 1 or abs(ch - state["last_ch"]) > 1:
                    # 如果没有手动缩放过，自动适应窗口大小 (响应式)
                    if not state.get("manual_zoom", False):
                        w, h = state["img"].size
                        if w > 0 and h > 0:
                             # 重新计算适应比例
                            new_scale = min(cw/w, ch/h)
                            new_scale = min(new_scale, 1.0) # 限制最大 100%
                            
                            # 如果比例变化明显，更新并重绘
                            if abs(new_scale - state["scale"]) > 0.001:
                                state["scale"] = new_scale
                                draw(force_resize=True)
                                return

                    draw(force_resize=False)

            canvas.bind("<MouseWheel>", on_wheel)
            canvas.bind("<ButtonPress-1>", start_pan)
            canvas.bind("<B1-Motion>", pan)
            canvas.bind("<Double-Button-1>", reset)
            canvas.bind("<Configure>", on_configure)
            
            CreateToolTip(canvas, "滚轮缩放，拖拽移动，双击复位")

        threading.Thread(target=_load_and_show, daemon=True).start()

    def _show_text_preview(self, container, path):
        """显示文本预览 (支持代码高亮、搜索、格式化、自动换行、跳转)"""
        # Grid 布局配置
        container.grid_columnconfigure(0, weight=0) # 行号列
        container.grid_columnconfigure(1, weight=1) # 内容列
        container.grid_rowconfigure(0, weight=0) # 工具栏
        container.grid_rowconfigure(1, weight=1) # 内容区域
        
        # === 0. 工具栏 ===
        toolbar = ttk.Frame(container, padding=2)
        toolbar.grid(row=0, column=0, columnspan=3, sticky="ew")
        
        # 自动换行切换
        wrap_var = tk.BooleanVar(value=False)
        
        def toggle_wrap():
            if wrap_var.get():
                txt.config(wrap="word")
                # 隐藏水平滚动条 (换行时不需要)
                xs.grid_remove()
            else:
                txt.config(wrap="none")
                xs.grid()
        
        ttk.Checkbutton(toolbar, text="自动换行", variable=wrap_var, command=toggle_wrap).pack(side="left", padx=5)
        
        # 跳转行
        def goto_line():
            try:
                # 简单输入框 (也可以用 simpledialog)
                top = tk.Toplevel(container)
                top.title("跳转到行")
                top.geometry("200x100")
                top.transient(container.winfo_toplevel())
                
                ttk.Label(top, text="行号:").pack(pady=5)
                e = ttk.Entry(top, width=10)
                e.pack(pady=5)
                e.focus_set()
                
                def do_goto(event=None):
                    val = e.get()
                    if val.isdigit():
                        line = int(val)
                        txt.see(f"{line}.0")
                        txt.mark_set("insert", f"{line}.0")
                        # 高亮一下
                        txt.tag_remove("goto_highlight", "1.0", "end")
                        txt.tag_add("goto_highlight", f"{line}.0", f"{line+1}.0")
                        txt.tag_config("goto_highlight", background=self.colors["listbox_sel_bg"], foreground=self.colors["listbox_sel_fg"])
                    top.destroy()
                
                e.bind("<Return>", do_goto)
                ttk.Button(top, text="跳转", command=do_goto).pack(pady=5)
            except:
                pass

        ttk.Button(toolbar, text="跳转到行...", command=goto_line, width=10).pack(side="left", padx=5)
        
        # 搜索按钮
        ttk.Button(toolbar, text="搜索 (Ctrl+F)", command=lambda: txt.event_generate("<Control-f>"), width=12).pack(side="left", padx=5)

        # 1. 行号区域
        ln_text = tk.Text(container, width=4, padx=4, takefocus=0, border=0,
                          background="#2b2b2b", foreground="#666666", state='disabled', wrap='none', font=("Consolas", 10))
        ln_text.grid(row=1, column=0, sticky='nsew')
        
        # 2. 内容区域
        txt = tk.Text(container, wrap="none", font=("Consolas", 10),
                      bg=self.colors["text_bg"],
                      fg=self.colors["text_fg"],
                      insertbackground=self.colors["fg"], # 光标颜色
                      highlightthickness=0,
                      borderwidth=0)
        txt.grid(row=1, column=1, sticky="nsew")
        
        # 3. 滚动条
        ys = ttk.Scrollbar(container, orient="vertical", command=lambda *args: (ln_text.yview(*args), txt.yview(*args)))
        xs = ttk.Scrollbar(container, orient="horizontal", command=txt.xview)
        ys.grid(row=1, column=2, sticky="ns")
        xs.grid(row=2, column=1, sticky="ew")
        
        # === 搜索栏 (Search Bar) ===
        search_frame = ttk.Frame(container, padding=2)
        search_var = tk.StringVar()
        entry = ttk.Entry(search_frame, textvariable=search_var, width=20)
        entry.pack(side="left", padx=2)
        
        lbl_match = ttk.Label(search_frame, text="0/0", font=("Segoe UI", 8))
        lbl_match.pack(side="left", padx=2)
        
        def find_text(direction=1):
            target = search_var.get()
            if not target: return
            
            start_pos = txt.index("insert")
            
            if direction == 1: # Next
                # search returns 'line.col'
                pos = txt.search(target, start_pos + "+1c", stopindex="end", nocase=True)
                if not pos: # Loop from start
                    pos = txt.search(target, "1.0", stopindex=start_pos, nocase=True)
            else: # Prev
                pos = txt.search(target, start_pos, stopindex="1.0", backwards=True, nocase=True)
                if not pos: # Loop from end
                    pos = txt.search(target, "end", stopindex=start_pos, backwards=True, nocase=True)
                    
            if pos:
                # Clear previous selection
                txt.tag_remove("search_highlight", "1.0", "end")
                txt.tag_config("search_highlight", background="#ffff00", foreground="#000000")
                
                end_pos = f"{pos}+{len(target)}c"
                txt.tag_add("search_highlight", pos, end_pos)
                txt.see(pos)
                txt.mark_set("insert", pos)
                lbl_match.config(text="Found")
            else:
                lbl_match.config(text="Not found")

        ttk.Button(search_frame, text="↓", width=2, command=lambda: find_text(1)).pack(side="left")
        ttk.Button(search_frame, text="↑", width=2, command=lambda: find_text(-1)).pack(side="left")
        ttk.Button(search_frame, text="×", width=2, command=lambda: (search_frame.grid_forget(), txt.focus_set())).pack(side="left", padx=5)

        entry.bind("<Return>", lambda e: find_text(1))
        entry.bind("<Shift-Return>", lambda e: find_text(-1))

        def toggle_search(e=None):
            if search_frame.winfo_viewable():
                search_frame.grid_forget()
                txt.focus_set()
            else:
                search_frame.grid(row=3, column=0, columnspan=3, sticky="ew")
                entry.focus_set()
            return "break"
        
        # 绑定 Ctrl+F
        txt.bind("<Control-f>", toggle_search)
        
        # 滚动联动
        def _on_scroll(*args):
            ys.set(*args)
            ln_text.yview_moveto(args[0])
        
        txt.configure(yscrollcommand=_on_scroll, xscrollcommand=xs.set)
        
        # 4. 读取与处理内容
        content = read_file_content(path)
        
        # JSON 格式化
        if path.lower().endswith('.json'):
            try:
                parsed = json.loads(content)
                content = json.dumps(parsed, indent=2, ensure_ascii=False)
                
                # 添加 "切换到树状视图" 按钮
                def switch_to_tree():
                    # 清空当前内容区域 (保留工具栏)
                    for widget in container.winfo_children():
                        if widget != toolbar:
                            widget.destroy()
                    
                    # 重新配置 Grid
                    container.grid_rowconfigure(1, weight=1)
                    
                    # 创建 Treeview
                    tree_frame = ttk.Frame(container)
                    tree_frame.grid(row=1, column=0, columnspan=3, sticky="nsew")
                    
                    tree = ttk.Treeview(tree_frame, columns=("value"), show="tree headings")
                    tree.heading("#0", text="Key")
                    tree.heading("value", text="Value")
                    
                    ys = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
                    xs = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
                    tree.configure(yscrollcommand=ys.set, xscrollcommand=xs.set)
                    
                    tree.pack(side="left", fill="both", expand=True)
                    ys.pack(side="right", fill="y")
                    xs.pack(side="bottom", fill="x")
                    
                    # 递归填充
                    def populate_tree(parent, data):
                        if isinstance(data, dict):
                            for k, v in data.items():
                                if isinstance(v, (dict, list)):
                                    node = tree.insert(parent, "end", text=k, open=False)
                                    populate_tree(node, v)
                                else:
                                    tree.insert(parent, "end", text=k, values=(str(v),))
                        elif isinstance(data, list):
                            for i, v in enumerate(data):
                                if isinstance(v, (dict, list)):
                                    node = tree.insert(parent, "end", text=f"[{i}]", open=False)
                                    populate_tree(node, v)
                                else:
                                    tree.insert(parent, "end", text=f"[{i}]", values=(str(v),))
                                    
                    populate_tree("", parsed)
                    
                    # 恢复按钮
                    btn_text.config(text="切换到文本视图", command=switch_to_text)

                def switch_to_text():
                    # 重新加载文本预览
                    for widget in container.winfo_children():
                        widget.destroy()
                    self._show_text_preview(container, path)

                btn_text = ttk.Button(toolbar, text="切换到树状视图", command=switch_to_tree)
                btn_text.pack(side="left", padx=5)
                
            except:
                pass
        
        # XML 格式化
        elif path.lower().endswith('.xml'):
            try:
                import xml.dom.minidom
                # Remove empty lines first to avoid messy output
                clean_xml = "".join([line.strip() for line in content.splitlines()])
                dom = xml.dom.minidom.parseString(content) # Use original content or clean? Original is safer.
                content = dom.toprettyxml(indent="  ")
                # Filter out excessive blank lines often caused by toprettyxml on already parsed xml
                content = "\n".join([line for line in content.splitlines() if line.strip()])
            except:
                pass
        
        # 性能优化：限制显示的行数
        MAX_LINES = 2000
        if content:
            lines = content.splitlines()
            if len(lines) > MAX_LINES:
                preview_content = "\n".join(lines[:MAX_LINES])
                preview_content += f"\n\n... (文件过长，仅显示前 {MAX_LINES} 行，共 {len(lines)} 行，请使用外部编辑器查看完整内容) ..."
                txt.insert("1.0", preview_content)
            else:
                txt.insert("1.0", content)
            
            # === 语法高亮 ===
            _, ext = os.path.splitext(path)
            self._highlight_syntax(txt, ext.lower())
        else:
            txt.insert("1.0", "")
        
        # 5. 生成行号
        line_count = int(txt.index('end-1c').split('.')[0])
        line_numbers_string = "\n".join(str(i) for i in range(1, line_count + 1))
        ln_text.config(state='normal')
        ln_text.insert('1.0', line_numbers_string)
        ln_text.config(state='disabled')
        
        txt.config(state="disabled")

    def _show_preview_column(self, name, full_path):
        """在最右侧显示预览面板"""
        # 计算初始宽度
        try:
            canvas_width = self.canvas.winfo_width()
            # 如果 canvas 还没显示出来，给个默认值
            if canvas_width <= 1: 
                canvas_width = 1000 
                
            used_width = 0
            for col in self.columns:
                if col["frame"].winfo_exists():
                    used_width += col["frame"].winfo_width() + 4
            
            initial_width = max(225, canvas_width - used_width - 4)
        except:
            initial_width = 225

        # 容器
        container = ttk.Frame(self.scroll_frame, width=initial_width)
        container.pack(side="left", fill="y", padx=2, pady=2)
        container.pack_propagate(False)
        
        # 标记这个容器属于预览列
        container.is_preview_container = True
        # 重置手动调整标志
        self._preview_manually_resized = False

        # === 修复预览列无法调整宽度的问题 ===
        # 在预览列左侧添加一个调整手柄 (Resize Handle)
        # 注意：为了更符合直觉，我们通常拖动右边界调整宽度，但对于最右侧的预览列，
        # 如果它填满了剩余空间，拖动左边界可能更合理？
        # 但按照目前 Miller Columns 的习惯，每一列都有右侧手柄。
        # 之前的代码已经在右侧添加了手柄，可能被覆盖或者不可见？
        
        # 内容
        frame = ttk.LabelFrame(container, text=f"预览: {name}", padding=2)
        frame.pack(side="left", fill="both", expand=True)

        # 拖拽手柄 (加宽并增加可见性，确保层级正确)
        # 注意：这里使用 side="right" pack，必须保证它在 frame 之前 pack 或者 frame 使用 expand=True
        # 在上面的代码中，frame 先 pack 且 expand=True，这会占据所有空间，导致 handle 被挤出或不可见
        # 修正：先 pack handle (side=right)，再 pack frame (side=left, expand=True)
        # 或者使用 grid 布局
        
        # 为了修复中间列调整失效的问题，我们统一调整 _add_column 和 _show_preview_column 的布局顺序
        # 但由于要最小化修改，我们这里先调整 pack 顺序：
        # 1. 移除 frame 的 pack
        frame.pack_forget()
        
        # 2. 创建并 pack 手柄
        handle = tk.Frame(container, width=14, bg="#333333", cursor="sb_h_double_arrow")
        handle.pack(side="right", fill="y")
        
        # 3. 重新 pack frame
        frame.pack(side="left", fill="both", expand=True)

        # 启用拖放 (预览列)
        if HAS_DND:
            frame.drop_target_register(DND_FILES)
            frame.dnd_bind('<<Drop>>', lambda e: self._on_drop_preview(e, full_path))
        
        # 手柄视觉反馈
        def _on_handle_enter(e): e.widget.config(bg="#555555")
        def _on_handle_leave(e): e.widget.config(bg="#333333")
        
        handle.bind("<Enter>", _on_handle_enter)
        handle.bind("<Leave>", _on_handle_leave)
        handle.bind("<Button-1>", lambda e: self._start_resize(e, container))
        handle.bind("<B1-Motion>", lambda e: self._perform_resize(e, container))
        
        # 记录为最后一列
        self.columns.append({
            "frame": container,
            "listbox": None, # 不是列表
            "path": full_path,
            "items": [],
            "is_preview": True
        })
        
        # 1. 工具栏 (移出 PanedWindow，直接放在 frame 顶部)
        toolbar = ttk.Frame(frame)
        toolbar.pack(fill="x", pady=(0,5))
        ttk.Button(toolbar, text="系统打开", command=lambda: self._open_system_file_path(full_path)).pack(side="right")
        
        # 上下分割
        paned = ttk.PanedWindow(frame, orient="vertical")
        paned.pack(fill="both", expand=True)
        
        # 2. 内容
        preview_area = ttk.Frame(paned)
        paned.add(preview_area, weight=3)
        
        if is_image_file(full_path):
            # === 图片预览 (新版：支持缩放/拖拽) ===
            self._show_image_preview(preview_area, full_path, initial_width)
        
        elif full_path.lower().endswith('.docx'):
            self._show_docx_preview(preview_area, full_path)
            
        elif is_video_file(full_path):
            # === 视频预览 (带控制) ===
            # 容器：上方是视频，下方是控制栏
            video_container = ttk.Frame(preview_area, style="Dark.TFrame")
            video_container.pack(fill="both", expand=True)
            
            # 1. 控制栏 (先 pack 确保底部空间保留)
            ctrl_frame = ttk.Frame(video_container, padding=5)
            ctrl_frame.pack(side="bottom", fill="x")

            # 2. 视频区域 (黑色背景，占据剩余空间)
            video_frame = tk.Frame(video_container, bg="black")
            video_frame.pack(side="top", fill="both", expand=True)
            
            label = ttk.Label(video_frame, anchor="center", background="black")
            label.pack(fill="both", expand=True)
            
            can_play_video = False
            try:
                import cv2
                can_play_video = True
            except ImportError:
                can_play_video = False

            # 检查音频支持 (移至线程中判断，避免主线程 import 卡顿)
            # can_play_audio = False
            # try:
            #     import ffpyplayer
            #     can_play_audio = True
            # except ImportError:
            #     pass

            if can_play_video and HAS_PIL:
                import queue
                # 优化：减小队列缓冲区 (10->3) 以减少延迟和内存占用
                q = queue.Queue(maxsize=3)
                stop_event = threading.Event()
                
                # IMPORTANT: Save stop_event to the column data so it can be triggered when column is cleared
                if self.columns and self.columns[-1]["path"] == full_path:
                    self.columns[-1]["video_stop_event"] = stop_event
                
                # 控制状态
                control_state = {
                    "paused": False,
                    "seek_req": None,
                    "dragging": False,
                    "duration": 0,
                    "current": 0,
                    "fps": 25, # 默认
                    "volume": 1.0,
                    "muted": False,
                    "view_width": initial_width, # 初始宽度
                    "view_height": 400 # 初始高度估计
                }
                
                # 监听容器大小变化
                def on_video_resize(event):
                    if event.width > 10:
                        control_state["view_width"] = event.width
                    if event.height > 10:
                        control_state["view_height"] = event.height
                
                video_frame.bind("<Configure>", on_video_resize)
                
                # 播放/暂停按钮
                def toggle_play(event=None):
                    control_state["paused"] = not control_state["paused"]
                    play_btn.config(text="▶" if control_state["paused"] else "⏸")
                    # 点击画面时获取焦点，以便接收键盘事件
                    if event: video_frame.focus_set()
                
                # 全屏相关逻辑
                def exit_fullscreen(event=None):
                    if "fs_win" in control_state:
                        try:
                            control_state["fs_win"].destroy()
                        except:
                            pass
                        del control_state["fs_win"]
                        del control_state["fs_label"]
                        # 恢复焦点到主窗口视频框
                        video_frame.focus_set()

                def toggle_fullscreen(event=None):
                    if "fs_win" in control_state:
                        exit_fullscreen()
                    else:
                        # 进入全屏
                        fs_win = tk.Toplevel(video_frame)
                        fs_win.title("全屏预览")
                        fs_win.attributes("-fullscreen", True)
                        fs_win.configure(background="black")
                        
                        # 全屏下的显示 Label
                        fs_label = ttk.Label(fs_win, anchor="center", background="black")
                        fs_label.pack(fill="both", expand=True)
                        
                        control_state["fs_win"] = fs_win
                        control_state["fs_label"] = fs_label
                        
                        # 绑定退出事件
                        fs_win.bind("<Escape>", exit_fullscreen)
                        fs_win.bind("<Double-Button-1>", exit_fullscreen)
                        
                        # 绑定控制快捷键 (转发给原来的处理函数)
                        fs_win.bind("<space>", lambda e: toggle_play())
                        fs_win.bind("<Left>", on_key_press)
                        fs_win.bind("<Right>", on_key_press)
                        fs_win.bind("<Up>", on_key_press)
                        fs_win.bind("<Down>", on_key_press)
                        
                        fs_win.focus_set()

                # 点击视频区域切换播放/暂停，双击全屏
                label.bind("<Button-1>", toggle_play)
                label.bind("<Double-Button-1>", toggle_fullscreen)
                CreateToolTip(label, "单击播放/暂停，双击全屏")
                
                play_btn = ttk.Button(ctrl_frame, text="⏸", width=3, command=toggle_play)
                play_btn.pack(side="left")
                CreateToolTip(play_btn, "播放/暂停 (Space)")
                
                # 截图按钮
                def take_snapshot():
                    if "current_image" in control_state:
                        try:
                            img = control_state["current_image"]
                            # 生成文件名
                            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                            base_name = os.path.splitext(os.path.basename(full_path))[0]
                            save_name = f"{base_name}_snapshot_{timestamp}.png"
                            save_path = os.path.join(os.path.dirname(full_path), save_name)
                            
                            img.save(save_path)
                            print(f"Snapshot saved: {save_path}")
                            # 简单的视觉反馈（如闪烁一下按钮）
                            orig_bg = snap_btn.cget("text")
                            snap_btn.config(text="✔")
                            self.after(1000, lambda: snap_btn.config(text="📷"))
                        except Exception as e:
                            print(f"Snapshot error: {e}")

                snap_btn = ttk.Button(ctrl_frame, text="📷", width=3, command=take_snapshot)
                snap_btn.pack(side="left", padx=2)
                CreateToolTip(snap_btn, "截图 (保存当前帧)")
                
                # 画质选择下拉框
                # Auto: 根据分辨率自动选择最佳帧率策略
                # 4K: 限制 20FPS (高性能模式)
                # 1080P: 限制 30FPS
                # 720P: 限制 30FPS + 降低分辨率 (提升性能)
                # 360P: 限制 30FPS + 大幅降低分辨率 (极致流畅)
                
                quality_var = tk.StringVar(value="Auto")
                
                def on_quality_change(event=None):
                    mode = quality_var.get()
                    control_state["quality_mode"] = mode
                    print(f"Quality changed to: {mode}")
                    # 重新聚焦视频以便快捷键可用
                    video_frame.focus_set()

                control_state["quality_mode"] = "Auto" # 默认

                quality_combo = ttk.Combobox(ctrl_frame, textvariable=quality_var, values=["Auto", "4K", "1080P", "720P", "360P"], width=6, state="readonly")
                quality_combo.pack(side="right", padx=2)
                quality_combo.bind("<<ComboboxSelected>>", on_quality_change)
                CreateToolTip(quality_combo, "画质选择 (影响清晰度和流畅度)")

                # 播放倍速选择
                speed_var = tk.StringVar(value="1.0x")
                
                def on_speed_change(event=None):
                    val_str = speed_var.get().replace("x", "")
                    try:
                        speed = float(val_str)
                        control_state["speed"] = speed
                        print(f"Speed changed to: {speed}")
                    except:
                        pass
                    # 重新聚焦视频
                    video_frame.focus_set()
                
                speed_combo = ttk.Combobox(ctrl_frame, textvariable=speed_var, values=["0.5x", "1.0x", "1.5x", "2.0x"], width=5, state="readonly")
                speed_combo.pack(side="right", padx=2)
                speed_combo.bind("<<ComboboxSelected>>", on_speed_change)
                CreateToolTip(speed_combo, "播放倍速 (注意: 非1.0x时可能会静音)")

                # 音量控制区域
                vol_frame = ttk.Frame(ctrl_frame)
                vol_frame.pack(side="right", padx=(5, 0))
                
                def toggle_mute():
                    control_state["muted"] = not control_state["muted"]
                    vol_btn.config(text="🔇" if control_state["muted"] else "🔊")
                    
                def on_vol_change(val):
                    control_state["volume"] = float(val)
                    if control_state["muted"]:
                        # 如果在静音状态下调节音量，自动取消静音
                        control_state["muted"] = False
                        vol_btn.config(text="🔊")
                
                vol_btn = ttk.Button(vol_frame, text="🔊", width=3, command=toggle_mute)
                vol_btn.pack(side="left")
                CreateToolTip(vol_btn, "静音/取消静音")
                
                vol_slider = ttk.Scale(vol_frame, from_=0.0, to=1.0, orient="horizontal", command=on_vol_change, length=60)
                vol_slider.set(1.0)
                vol_slider.pack(side="left", padx=2)
                CreateToolTip(vol_slider, "音量调节 (↑/↓)")
                
                # === 键盘快捷键控制 ===
                def on_key_press(event):
                    # 只有当视频区域获得焦点时才响应
                    if event.keysym == "space":
                        toggle_play()
                    elif event.keysym == "Left":
                        # 快退 5秒
                        cur = control_state.get('current', 0)
                        control_state["seek_req"] = max(0, cur - 5)
                    elif event.keysym == "Right":
                        # 快进 5秒
                        cur = control_state.get('current', 0)
                        dur = control_state.get('duration', 0)
                        control_state["seek_req"] = min(dur, cur + 5)
                    elif event.keysym == "Up":
                        # 音量 +10%
                        vol = control_state.get('volume', 1.0)
                        new_vol = min(1.0, vol + 0.1)
                        vol_slider.set(new_vol)
                        on_vol_change(new_vol)
                    elif event.keysym == "Down":
                        # 音量 -10%
                        vol = control_state.get('volume', 1.0)
                        new_vol = max(0.0, vol - 0.1)
                        vol_slider.set(new_vol)
                        on_vol_change(new_vol)
                        
                video_frame.bind("<Key>", on_key_press)
                # 确保 label 点击也能把焦点传给 video_frame
                label.bind("<Button-1>", lambda e: (video_frame.focus_set(), toggle_play()))

                # 时间显示 (稍微加宽以容纳帧信息)
                time_lbl = ttk.Label(ctrl_frame, text="00:00:00 / 00:00:00", width=18, anchor="center")
                time_lbl.pack(side="right")
                
                # 进度条回调
                def on_slider_change(val):
                    fval = float(val)
                    # 仅在拖拽时处理，避免播放时自我触发
                    if control_state.get("dragging", False):
                        control_state["seek_req"] = fval
                        # 拖拽时实时更新时间标签
                        d = control_state.get('duration', 0)
                        fps = control_state.get('fps', 25)
                        time_lbl.config(text=f"{self._format_time_frames(fval, fps)} / {self._format_time_frames(d, fps)}")
                    
                slider = ttk.Scale(ctrl_frame, from_=0, to=100, orient="horizontal", command=on_slider_change)
                slider.pack(side="left", fill="x", expand=True, padx=5)
                
                # === 进度条交互优化：支持点击跳转 ===
                def update_seek_from_event(e):
                    """根据鼠标点击位置计算进度"""
                    width = slider.winfo_width()
                    if width > 5:
                        ratio = e.x / width
                        ratio = max(0.0, min(1.0, ratio))
                        duration = slider.cget('to')
                        new_val = ratio * duration
                        
                        slider.set(new_val)
                        # 手动触发更新逻辑
                        on_slider_change(new_val)

                def on_seek_start(e):
                    control_state["dragging"] = True
                    update_seek_from_event(e)
                    return "break" # 接管事件，防止默认行为冲突

                def on_seek_motion(e):
                    if control_state.get("dragging", False):
                        update_seek_from_event(e)
                    return "break"

                def on_seek_end(e):
                    control_state["dragging"] = False
                    return "break"

                slider.bind("<ButtonPress-1>", on_seek_start)
                slider.bind("<B1-Motion>", on_seek_motion)
                slider.bind("<ButtonRelease-1>", on_seek_end)
                
                # 初始设置 Range 需要等到 duration 获取后 (在 update 中动态调整)
                # 这里先设置一个默认最大值
                slider.config(to=100) 

                self.columns[-1]["video_stop_event"] = stop_event
                
                # 保存控制函数供全局调用
                self.columns[-1]["video_control"] = {
                    "toggle_play": toggle_play
                }
                
                # 信息标签 (显示在控制栏上方)
                info_lbl = ttk.Label(video_container, text="", font=("Segoe UI", 8), anchor="center", foreground="#888888")
                info_lbl.pack(side="bottom", fill="x", pady=(0, 2))
                
                # 右键菜单
                menu = tk.Menu(video_frame, tearoff=0)
                
                # 倍速子菜单
                speed_menu = tk.Menu(menu, tearoff=0)
                control_state["speed"] = 1.0
                
                def set_speed(s):
                    control_state["speed"] = s
                    # 如果倍速不为 1.0，自动静音以防声画不同步
                    if s != 1.0:
                        control_state["muted"] = True
                        vol_btn.config(text="🔇")
                    
                speed_menu.add_radiobutton(label="0.5x", command=lambda: set_speed(0.5))
                speed_menu.add_radiobutton(label="1.0x", command=lambda: set_speed(1.0))
                speed_menu.add_radiobutton(label="1.5x", command=lambda: set_speed(1.5))
                speed_menu.add_radiobutton(label="2.0x", command=lambda: set_speed(2.0))
                # 默认选中 1.0x
                speed_menu.invoke(1)
                
                menu.add_cascade(label="播放速度", menu=speed_menu)
                menu.add_separator()
                menu.add_command(label="全屏 (Double Click)", command=toggle_fullscreen)
                menu.add_command(label="截图 (Snapshot)", command=take_snapshot)
                
                def show_context_menu(event):
                    menu.post(event.x_root, event.y_root)
                    
                label.bind("<Button-3>", show_context_menu)
                
                threading.Thread(target=self._video_thread, args=(full_path, q, stop_event, control_state), daemon=True).start()
                self._update_video_label(label, q, stop_event, control_state, slider, time_lbl, play_btn, info_lbl)
                
                # 监听 duration 变化以更新 slider max
                def check_duration():
                    if stop_event.is_set(): return
                    d = control_state.get('duration', 0)
                    if d > 0 and slider.cget('to') != d:
                        slider.config(to=d)
                    self.after(500, check_duration)
                check_duration()
                
            else:
                if not can_play_video:
                    msg = "视频预览需要 opencv-python\n请 pip install opencv-python"
                else:
                    msg = "需要 PIL"
                
                # if can_play_video and not can_play_audio:
                #     msg = "当前支持无声播放\n安装 ffpyplayer 可支持声音"
                    
                label.config(text=msg, foreground="white")

        elif full_path.lower().endswith(('.csv', '.xlsx', '.xls', '.et')):
            # === 表格预览 (CSV/Excel) ===
            if full_path.lower().endswith('.csv'):
                self._show_csv_preview(preview_area, full_path)
            else:
                self._show_xlsx_preview(preview_area, full_path)

        elif full_path.lower().endswith(('.docx', '.doc')):
            # === Word 文档预览 ===
            self._show_docx_preview(preview_area, full_path)

        elif full_path.lower().endswith('.pdf'):
            # === PDF 文档预览 ===
            self._show_pdf_preview(preview_area, full_path)

        elif full_path.lower().endswith(('.zip', '.tar', '.tar.gz', '.tgz', '.jar')):
            # === 压缩包预览 ===
            self._show_archive_preview(preview_area, full_path)

        elif is_text_file(full_path):
            # === 文本预览 (新版：支持 XML/JSON 格式化) ===
            self._show_text_preview(preview_area, full_path)
        else:
            # === 二进制/未知格式预览 (Hex View) ===
            self._show_hex_preview(preview_area, full_path)
        
        # 3. 信息
        info_frame = ttk.LabelFrame(paned, text="信息", padding=5, height=150)
        info_frame.pack_propagate(False) # 固定高度
        paned.add(info_frame, weight=0)  # weight=0 表示不参与自动伸缩
        
        info_txt = tk.Text(info_frame, wrap="word", state="disabled", 
                           bg=self.colors["text_bg"], 
                           fg=self.colors["text_fg"],
                           highlightthickness=0,
                           borderwidth=0)
        info_txt.pack(fill="both", expand=True)
        
        try:
            stat = os.stat(full_path)
            size_kb = stat.st_size / 1024
            created = datetime.datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            modified = datetime.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            
            info_str = f"路径: {full_path}\n大小: {size_kb:.2f} KB\n修改: {modified}\n创建: {created}"
            
            info_txt.config(state="normal")
            info_txt.insert("1.0", info_str)
            info_txt.config(state="disabled")
            
            # 计算 Hash (异步执行，避免卡顿)
            if stat.st_size < 50 * 1024 * 1024: # < 50MB (异步可以放宽限制)
                def _calc_md5():
                    try:
                        import hashlib
                        with open(full_path, "rb") as f:
                            data = f.read()
                            md5 = hashlib.md5(data).hexdigest()
                        
                        def _update_info():
                            if info_txt.winfo_exists():
                                info_txt.config(state="normal")
                                info_txt.insert("end", f"\nMD5: {md5}")
                                info_txt.config(state="disabled")
                        self.after(0, _update_info)
                    except: pass
                
                threading.Thread(target=_calc_md5, daemon=True).start()
        except:
            pass
            
        # 自动滚动
        self.canvas.update_idletasks()
        self.canvas.xview_moveto(1.0)

    def _show_docx_preview(self, container, path):
        """显示 Word 文档预览 (.docx)"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        # 使用 Text 组件显示内容
        txt = tk.Text(container, wrap="word", font=("Segoe UI", 11),
                      bg=self.colors["text_bg"],
                      fg=self.colors["text_fg"],
                      insertbackground=self.colors["fg"],
                      highlightthickness=0,
                      borderwidth=0)
        
        ys = ttk.Scrollbar(container, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        
        txt.grid(row=0, column=0, sticky="nsew")
        ys.grid(row=0, column=1, sticky="ns")
        
        try:
            import docx
            doc = docx.Document(path)
            full_text = []
            
            # 读取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # 读取表格 (简单追加在后面)
            if doc.tables:
                full_text.append("\n--- 表格内容 ---\n")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        full_text.append(" | ".join(row_text))
                    full_text.append("") # 空行分隔表格
            
            content = "\n\n".join(full_text)
            
            # 限制显示长度
            MAX_CHARS = 50000
            if len(content) > MAX_CHARS:
                content = content[:MAX_CHARS] + "\n\n... (文档过长，仅显示前 50000 字符) ..."
            
            if not content.strip():
                content = "(文档为空)"

            txt.insert("1.0", content)
            
        except ImportError:
            txt.insert("1.0", "需要安装 python-docx 库才能预览 Word 文档。\n请运行: pip install python-docx")
        except Exception as e:
            # 尝试处理 .doc (python-docx 不支持 .doc，通常需要 win32com 或转换)
            if path.lower().endswith('.doc'):
                txt.insert("1.0", f"不支持直接预览 .doc 格式 (仅支持 .docx)。\n建议另存为 .docx 后查看。\n\n错误信息: {e}")
            else:
                txt.insert("1.0", f"读取文档失败: {e}")
            
        txt.config(state="disabled")

    def _show_docx_preview(self, container, path):
        """显示 DOCX 文档预览"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        txt = tk.Text(container, wrap="word", font=("Georgia", 11),
                      bg=self.colors["text_bg"], fg=self.colors["text_fg"],
                      highlightthickness=0, borderwidth=0)
        
        ys = ttk.Scrollbar(container, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        
        txt.grid(row=0, column=0, sticky="nsew")
        ys.grid(row=0, column=1, sticky="ns")
        
        try:
            import docx
            doc = docx.Document(path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            content = "\n\n".join(full_text)
            
            # 限制显示长度
            if len(content) > 50000:
                content = content[:50000] + "\n\n... (文档过长，已截断) ..."
                
            txt.insert("1.0", content)
            txt.config(state="disabled")
            
        except ImportError:
            txt.insert("1.0", "预览失败: 需安装 python-docx 库\n请运行: pip install python-docx")
        except Exception as e:
            txt.insert("1.0", f"读取 DOCX 失败: {e}")

    def _show_pdf_preview(self, container, path):
        """显示 PDF 文档预览 (仅提取文本)"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        txt = tk.Text(container, wrap="word", font=("Segoe UI", 11),
                      bg=self.colors["text_bg"],
                      fg=self.colors["text_fg"],
                      insertbackground=self.colors["fg"],
                      highlightthickness=0,
                      borderwidth=0)
        
        ys = ttk.Scrollbar(container, orient="vertical", command=txt.yview)
        txt.configure(yscrollcommand=ys.set)
        
        txt.grid(row=0, column=0, sticky="nsew")
        ys.grid(row=0, column=1, sticky="ns")
        
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            
            # 元数据
            meta_info = []
            if reader.metadata:
                if reader.metadata.title: meta_info.append(f"标题: {reader.metadata.title}")
                if reader.metadata.author: meta_info.append(f"作者: {reader.metadata.author}")
            
            meta_info.append(f"总页数: {len(reader.pages)}")
            txt.insert("1.0", "\n".join(meta_info) + "\n\n" + "-"*30 + "\n\n")
            
            # 读取内容 (限制前 20 页以防卡顿)
            MAX_PAGES = 20
            full_text = []
            
            for i, page in enumerate(reader.pages):
                if i >= MAX_PAGES:
                    full_text.append(f"\n... (仅显示前 {MAX_PAGES} 页) ...")
                    break
                
                page_text = page.extract_text()
                if page_text:
                    full_text.append(f"--- 第 {i+1} 页 ---\n{page_text}\n")
            
            content = "\n".join(full_text)
            if not content.strip():
                content = "(PDF 内容为空或无法提取文本)"
                
            txt.insert("end", content)
            
        except ImportError:
            txt.insert("1.0", "需要安装 pypdf 库才能预览 PDF 文档。\n请运行: pip install pypdf")
        except Exception as e:
            txt.insert("1.0", f"读取 PDF 失败: {e}")
            
        txt.config(state="disabled")

    def _show_archive_preview(self, container, path):
        """显示压缩包内容列表 (.zip, .tar)"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        columns = ("文件路径", "大小", "修改时间")
        
        # 滚动条
        tree_scroll_y = ttk.Scrollbar(container)
        tree_scroll_x = ttk.Scrollbar(container, orient="horizontal")
        
        tree = ttk.Treeview(container, columns=columns, show="headings",
                            yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        tree.grid(row=0, column=0, sticky="nsew")
        tree_scroll_y.grid(row=0, column=1, sticky="ns")
        tree_scroll_x.grid(row=1, column=0, sticky="ew")
        
        # 表头
        tree.heading("文件路径", text="文件路径")
        tree.heading("大小", text="大小")
        tree.heading("修改时间", text="修改时间")
        
        tree.column("文件路径", width=300, minwidth=100)
        tree.column("大小", width=100, minwidth=50)
        tree.column("修改时间", width=150, minwidth=100)
        
        # 样式适配
        style = ttk.Style()
        style.configure("Treeview", 
            background=self.colors["text_bg"], 
            foreground=self.colors["text_fg"], 
            fieldbackground=self.colors["text_bg"],
            borderwidth=0
        )
        
        try:
            import zipfile
            import tarfile
            import datetime
            
            items = []
            
            if zipfile.is_zipfile(path):
                with zipfile.ZipFile(path, 'r') as zf:
                    for info in zf.infolist():
                        dt = datetime.datetime(*info.date_time).strftime('%Y-%m-%d %H:%M:%S')
                        size_str = f"{info.file_size / 1024:.1f} KB"
                        items.append((info.filename, size_str, dt))
                        
            elif tarfile.is_tarfile(path):
                with tarfile.open(path, 'r') as tf:
                    for member in tf.getmembers():
                        dt = datetime.datetime.fromtimestamp(member.mtime).strftime('%Y-%m-%d %H:%M:%S')
                        size_str = f"{member.size / 1024:.1f} KB"
                        items.append((member.name, size_str, dt))
            
            # 插入数据 (限制 500 条)
            for i, item in enumerate(items):
                if i >= 500:
                    tree.insert("", "end", values=("... (仅显示前 500 个文件) ...", "", ""))
                    break
                tree.insert("", "end", values=item)
                
            if not items:
                tree.insert("", "end", values=("(压缩包为空)", "", ""))
                
        except Exception as e:
            tk.Label(container, text=f"读取压缩包失败: {e}", fg="red").grid(row=0, column=0)
            return

    def _show_csv_preview(self, container, path):
        """显示 CSV 表格预览"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        # 提示标签
        try:
            # 预读取部分行以确定列
            rows = []
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                # 读取前 200 行
                for i, row in enumerate(reader):
                    if i >= 200: break
                    rows.append(row)
        except Exception as e:
            tk.Label(container, text=f"CSV 读取错误: {e}", fg="red").grid(row=0, column=0)
            return

        if not rows:
            tk.Label(container, text="CSV 文件为空", fg="gray").grid(row=0, column=0)
            return
            
        columns = rows[0]
        data = rows[1:]

        # 滚动条
        tree_scroll_y = ttk.Scrollbar(container)
        tree_scroll_x = ttk.Scrollbar(container, orient="horizontal")
        
        # 配置 Treeview 样式 (跟随主题)
        style = ttk.Style()
        style.configure("Treeview", background=self.colors["text_bg"], foreground=self.colors["text_fg"], fieldbackground=self.colors["text_bg"])
        style.configure("Treeview.Heading", background=self.colors["listbox_sel_bg"], foreground=self.colors["listbox_sel_fg"])
        style.map("Treeview", background=[("selected", self.colors["listbox_sel_bg"])], foreground=[("selected", self.colors["listbox_sel_fg"])])

        # 配置 Treeview 样式 (跟随主题)
        style = ttk.Style()
        style.configure("Treeview", background=self.colors["text_bg"], foreground=self.colors["text_fg"], fieldbackground=self.colors["text_bg"])
        style.configure("Treeview.Heading", background=self.colors["listbox_sel_bg"], foreground=self.colors["listbox_sel_fg"])
        style.map("Treeview", background=[("selected", self.colors["listbox_sel_bg"])], foreground=[("selected", self.colors["listbox_sel_fg"])])

        # 创建 Treeview
        tree = ttk.Treeview(container, columns=columns, show="headings", 
                            yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        tree.grid(row=0, column=0, sticky="nsew")
        tree_scroll_y.grid(row=0, column=1, sticky="ns")
        tree_scroll_x.grid(row=1, column=0, sticky="ew")
        
        # 设置表头和列宽
        import tkinter.font as tkfont
        font = tkfont.Font()
        
        for col_idx, col in enumerate(columns):
            col_str = str(col)
            tree.heading(col_str, text=col_str)
            
            # 计算最佳宽度
            # 1. 表头宽度
            max_width = font.measure(col_str) + 20
            
            # 2. 内容宽度 (采样前 100 行)
            for row in data[:100]:
                if col_idx < len(row):
                    cell_width = font.measure(str(row[col_idx])) + 20
                    max_width = max(max_width, cell_width)
            
            # 3. 限制范围
            final_width = min(400, max(50, max_width))
            tree.column(col_str, width=final_width, minwidth=50)
            
        # 插入数据
        for row in data:
            # 补齐 row 长度以免报错
            values = list(row)
            if len(values) < len(columns):
                values += [""] * (len(columns) - len(values))
            elif len(values) > len(columns):
                values = values[:len(columns)]
            tree.insert("", "end", values=values)
            
        if len(rows) >= 200:
            lbl = tk.Label(container, text="提示: 为保证性能，仅显示前 200 行", fg="#888888", bg="#2b2b2b")
            lbl.grid(row=2, column=0, sticky="w", padx=5)

    def _show_xlsx_preview(self, container, path):
        """显示 Excel/WPS 表格预览 (.xlsx, .xls, .et)"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        rows = []
        error_msg = None
        
        try:
            # 优先尝试 openpyxl (支持 .xlsx, .xlsm, .xltx, .xltm)
            # WPS 的 .et 实际上经常是兼容格式，也可以尝试
            import openpyxl
            try:
                wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                ws = wb.active
                # 读取前 200 行
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 200: break
                    rows.append(list(row))
                wb.close()
            except Exception as e_openpyxl:
                # 如果 openpyxl 失败 (例如 .xls 格式)，尝试 xlrd
                error_msg = f"openpyxl failed: {e_openpyxl}"
                try:
                    import xlrd
                    wb = xlrd.open_workbook(path)
                    sheet = wb.sheet_by_index(0)
                    for i in range(min(200, sheet.nrows)):
                        rows.append(sheet.row_values(i))
                    error_msg = None # 成功修复
                except Exception as e_xlrd:
                    if error_msg:
                        error_msg += f"\nxlrd failed: {e_xlrd}"
                    else:
                        error_msg = f"xlrd failed: {e_xlrd}"

        except ImportError:
            # 尝试捕获 import 错误，分别提示
            missing = []
            try: import openpyxl
            except ImportError: missing.append("openpyxl")
            try: import xlrd
            except ImportError: missing.append("xlrd")
            
            error_msg = f"缺少依赖库: {', '.join(missing)}\n请运行: pip install {' '.join(missing)}"
            
        except Exception as e:
            error_msg = f"未知错误: {e}"

        if error_msg and not rows:
            tk.Label(container, text=f"无法读取 Excel 文件:\n{error_msg}", fg="red", justify="left").grid(row=0, column=0)
            return

        if not rows:
            tk.Label(container, text="Excel 文件为空", fg="gray").grid(row=0, column=0)
            return
            
        # 处理表头 (第一行)
        columns = [str(col) if col is not None else "" for col in rows[0]]
        # 处理数据 (后续行)
        data = []
        for r in rows[1:]:
             data.append([str(cell) if cell is not None else "" for cell in r])

        # 滚动条
        tree_scroll_y = ttk.Scrollbar(container)
        tree_scroll_x = ttk.Scrollbar(container, orient="horizontal")
        
        # 创建 Treeview
        tree = ttk.Treeview(container, columns=columns, show="headings", 
                            yscrollcommand=tree_scroll_y.set, xscrollcommand=tree_scroll_x.set)
        
        tree_scroll_y.config(command=tree.yview)
        tree_scroll_x.config(command=tree.xview)
        
        tree.grid(row=0, column=0, sticky="nsew")
        tree_scroll_y.grid(row=0, column=1, sticky="ns")
        tree_scroll_x.grid(row=1, column=0, sticky="ew")
        
        # 设置表头和列宽
        import tkinter.font as tkfont
        font = tkfont.Font()

        for col_idx, col in enumerate(columns):
            col_str = str(col)
            tree.heading(col_str, text=col_str)
            
            # 计算最佳宽度
            # 1. 表头宽度
            max_width = font.measure(col_str) + 20
            
            # 2. 内容宽度 (采样前 100 行)
            for row in data[:100]:
                if col_idx < len(row):
                    cell_width = font.measure(str(row[col_idx])) + 20
                    max_width = max(max_width, cell_width)
            
            # 3. 限制范围
            final_width = min(400, max(50, max_width))
            tree.column(col_str, width=final_width, minwidth=50)
            
        # 插入数据
        for row in data:
            # 补齐 row 长度以免报错
            values = list(row)
            if len(values) < len(columns):
                values += [""] * (len(columns) - len(values))
            elif len(values) > len(columns):
                values = values[:len(columns)]
            tree.insert("", "end", values=values)
            
        if len(rows) >= 200:
            lbl = tk.Label(container, text="提示: 为保证性能，仅显示前 200 行", fg="#888888", bg="#2b2b2b")
            lbl.grid(row=2, column=0, sticky="w", padx=5)

    def _show_hex_preview(self, container, path):
        """显示十六进制预览"""
        # Grid 布局
        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)
        
        txt = tk.Text(container, wrap="none", font=("Courier New", 10),
                      bg=self.colors["text_bg"], fg=self.colors["text_fg"],
                      highlightthickness=0, borderwidth=0)
        
        ys = ttk.Scrollbar(container, orient="vertical", command=txt.yview)
        xs = ttk.Scrollbar(container, orient="horizontal", command=txt.xview)
        txt.configure(yscrollcommand=ys.set, xscrollcommand=xs.set)
        
        txt.grid(row=0, column=0, sticky="nsew")
        ys.grid(row=0, column=1, sticky="ns")
        xs.grid(row=1, column=0, sticky="ew")
        
        try:
            with open(path, "rb") as f:
                content = f.read(2048) # 只读取前 2KB
                
            hex_output = []
            for i in range(0, len(content), 16):
                chunk = content[i:i+16]
                
                # 地址
                addr = f"{i:08x}"
                
                # Hex 部分
                hex_vals = " ".join(f"{b:02x}" for b in chunk)
                padding = "   " * (16 - len(chunk))
                
                # ASCII 部分
                ascii_vals = "".join((chr(b) if 32 <= b < 127 else ".") for b in chunk)
                
                hex_output.append(f"{addr}  {hex_vals}{padding}  |{ascii_vals}|")
            
            if len(content) == 2048:
                hex_output.append("\n... (只显示前 2KB) ...")
                
            txt.insert("1.0", "\n".join(hex_output))
            txt.config(state="disabled")
            
        except Exception as e:
            txt.insert("1.0", f"读取失败: {e}")

    def _open_system_file_path(self, full_path):
        try:
            if hasattr(os, "startfile"):
                os.startfile(full_path)
            else:
                # import subprocess, sys # 已在顶部导入
                opener = "open" if sys.platform == "darwin" else "xdg-open"
                subprocess.call([opener, full_path])
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文件:\n{e}")

    def _on_right_click(self, event, col_index):
        """右键菜单 (支持空白处右键)"""
        lb = event.widget
        # 确保右键点击时获得焦点，以便粘贴等操作作用于当前列
        lb.focus_set()
        
        try:
            # 1. 判断点击位置
            index = lb.nearest(event.y)
            bbox = lb.bbox(index)
            
            clicked_on_item = False
            if bbox and event.y <= bbox[1] + bbox[3]:
                clicked_on_item = True
            
            if clicked_on_item:
                # 智能选择逻辑：
                # 如果点击的项不在当前选区中，则单选该项
                # 如果点击的项已经在选区中（多选状态下），则保持选区不变
                sel = lb.curselection()
                if index not in sel:
                    lb.selection_clear(0, tk.END)
                    lb.selection_set(index)
                    lb.activate(index)
                    # 触发选中逻辑
                    self._on_column_select(col_index)
                
                # 2. 获取当前列信息
                col_data = self.columns[col_index]
                current_dir = col_data["path"]
                
                # 3. 创建菜单
                m = tk.Menu(self, tearoff=0)
                
                # === 针对文件的菜单项 ===
                # 获取文件名等信息 (取第一个选中的作为主操作对象，或者根据命令处理全部)
                if index < len(col_data["filtered_items"]):
                    name, full_path, is_dir, _, _ = col_data["filtered_items"][index]
                    
                    # 仅当单选时才显示"打开所在文件夹" (避免歧义)
                    if len(lb.curselection()) == 1:
                        m.add_command(label="📂 打开所在文件夹", command=lambda: self._open_file_location(col_index, index))
                        m.add_separator()
                    
                    m.add_command(label="📋 复制 (Ctrl+C)", command=lambda: self._perform_copy(col_index))
                    m.add_command(label="✂️ 剪切", command=lambda: self._perform_cut(col_index))
                    m.add_command(label="📋 粘贴 (Ctrl+V)", command=lambda: self._perform_paste(col_index))
                    m.add_separator()
                    m.add_command(label="✏️ 重命名", command=lambda: self._rename_item(col_index, index))
                    m.add_command(label="🗑️ 删除 (Delete)", command=lambda: self._perform_delete(col_index))
                    m.add_separator()
                    
                    if len(lb.curselection()) == 1:
                        m.add_command(label="复制完整路径", command=lambda: self.clipboard_clear() or self.clipboard_append(full_path))
                        m.add_command(label="复制文件名", command=lambda: self.clipboard_clear() or self.clipboard_append(name))
            else:
                # 点击在空白处，清除选择
                lb.selection_clear(0, tk.END)
                
                # 2. 获取当前列信息
                col_data = self.columns[col_index]
                current_dir = col_data["path"]
                
                # 3. 创建菜单
                m = tk.Menu(self, tearoff=0)
                
                # === 针对空白处的菜单项 (文件夹操作) ===
                m.add_command(label="📋 粘贴 (Ctrl+V)", command=lambda: self._perform_paste(col_index))
                m.add_separator()
                m.add_command(label="➕ 新建文件夹", command=lambda: self._new_folder(col_index))
                m.add_command(label="🔄 刷新", command=lambda: self._reload_column(col_index))
            
            # 公共菜单项
            m.add_separator()
            m.add_command(label="💻 在终端打开", command=lambda: self._open_terminal(current_dir))
            m.add_command(label="ℹ️ 属性", command=lambda: self._show_properties(col_index, index if clicked_on_item else None))
            
            m.tk_popup(event.x_root, event.y_root)
        except Exception as e:
            print(f"Right click error: {e}")

    def _show_properties(self, col_index, item_index):
        col_data = self.columns[col_index]
        if item_index is not None and item_index < len(col_data["filtered_items"]):
            name, path, is_dir, size, mtime = col_data["filtered_items"][item_index]
        else:
            path = col_data["path"]
            name = os.path.basename(path)
            is_dir = True
            
        try:
            stat = os.stat(path)
            size_mb = stat.st_size / (1024 * 1024)
            ctime = datetime.datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S')
            mtime = datetime.datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')
            
            msg = f"名称: {name}\n" \
                  f"类型: {'文件夹' if is_dir else '文件'}\n" \
                  f"位置: {path}\n" \
                  f"大小: {size_mb:.2f} MB ({stat.st_size} bytes)\n" \
                  f"创建时间: {ctime}\n" \
                  f"修改时间: {mtime}"
                  
            messagebox.showinfo("属性", msg)
        except Exception as e:
            messagebox.showerror("错误", f"无法获取属性: {e}")

            # 启用/禁用粘贴
            # 始终启用粘贴，点击后再检查剪贴板，避免误判导致无法粘贴
            # has_files_to_paste = ... (移除复杂的预检查逻辑)
            pass
            
            # 排序子菜单 (仅在有内容时显示，或者始终显示)
            if col_data["all_items"]:
                sort_menu = tk.Menu(m, tearoff=0)
                sort_menu.add_command(label="按名称", command=lambda: self._sort_column(col_index, 'name'))
                sort_menu.add_command(label="按修改时间", command=lambda: self._sort_column(col_index, 'date'))
                sort_menu.add_command(label="按大小", command=lambda: self._sort_column(col_index, 'size'))
                m.add_cascade(label="排序方式", menu=sort_menu)
            
            m.post(event.x_root, event.y_root)
        except Exception as e:
            print(f"Right click error: {e}")

    def _on_enter_key(self, col_index):
        """回车键处理"""
        col_data = self.columns[col_index]
        sel = col_data["listbox"].curselection()
        if not sel: return
        
        index = sel[0]
        name, full_path, is_dir, _, _ = col_data["filtered_items"][index]
        
        if not is_dir:
            # 文件：系统打开
            self._open_system_file_path(full_path)
        else:
            # 文件夹：确保下一列获得焦点（如果有）
            if col_index + 1 < len(self.columns):
                next_col = self.columns[col_index + 1]
                if next_col.get("listbox"):
                    next_col["listbox"].focus_set()
                    if next_col["listbox"].size() > 0:
                        next_col["listbox"].selection_clear(0, tk.END)
                        next_col["listbox"].selection_set(0)
                        next_col["listbox"].activate(0)
                        self._on_column_select(col_index + 1)

    def _open_terminal(self, path):
        """在指定路径打开终端"""
        try:
            if sys.platform == "win32":
                subprocess.Popen(f'start cmd /k "cd /d {path}"', shell=True)
            elif sys.platform == "darwin":
                subprocess.run(['open', '-a', 'Terminal', path])
            else:
                # 尝试 gnome-terminal 或 xterm
                try:
                    subprocess.Popen(['gnome-terminal', '--working-directory', path])
                except:
                    subprocess.Popen(['xterm', '-e', f'cd "{path}" && /bin/bash'])
        except Exception as e:
            messagebox.showerror("错误", f"无法打开终端:\n{e}")

    def _get_file_icon(self, name, is_dir):
        """获取文件图标"""
        if is_dir:
            return "📂 "
        
        _, ext = os.path.splitext(name)
        ext = ext.lower()
        
        icons = {
            # 代码
            ".py": "🐍 ", ".js": "⚡ ", ".ts": "📘 ", ".jsx": "⚛️ ", ".tsx": "⚛️ ",
            ".html": "🌐 ", ".css": "🎨 ", ".scss": "🎨 ", ".less": "🎨 ",
            ".json": "🔧 ", ".xml": "⚙️ ", ".yaml": "⚙️ ", ".yml": "⚙️ ",
            ".sql": "🗄️ ", ".java": "☕ ", ".c": "🇨 ", ".cpp": "🇨 ", ".go": "🐹 ",
            # 文档
            ".md": "📝 ", ".txt": "📄 ", ".log": "📋 ", ".pdf": "📕 ",
            ".doc": "📘 ", ".docx": "📘 ", ".xls": "📗 ", ".xlsx": "📗 ",
            # 媒体
            ".png": "🖼️ ", ".jpg": "🖼️ ", ".jpeg": "🖼️ ", ".gif": "🎞️ ", ".webp": "🖼️ ",
            ".mp3": "🎵 ", ".wav": "🎵 ", ".mp4": "🎬 ", ".mov": "🎬 ",
            # 压缩包
            ".zip": "📦 ", ".rar": "📦 ", ".7z": "📦 ", ".tar": "📦 ", ".gz": "📦 ",
            # 系统
            ".exe": "💾 ", ".bat": "💻 ", ".sh": "🐚 ", ".ps1": "💻 ",
            ".dll": "⚙️ ", ".ini": "⚙️ "
        }
        return icons.get(ext, "📄 ")

    def _filter_items(self, col_index, search_text):
        """根据搜索文本过滤列表 (带数量限制优化)"""
        if col_index >= len(self.columns): return
        col_data = self.columns[col_index]
        all_items = col_data["all_items"]
        lb = col_data["listbox"]
        
        lb.delete(0, tk.END)
        filtered = []
        
        dir_count = 0
        file_count = 0
        
        MAX_DISPLAY = 2000 # 最大显示数量，防止卡死
        
        for item in all_items:
            name, full_path, is_dir, _, _ = item
            if not search_text or search_text in name.lower():
                # 仅当未达到显示限制时才插入 Listbox
                if len(filtered) < MAX_DISPLAY:
                    prefix = self._get_file_icon(name, is_dir)
                    lb.insert(tk.END, prefix + name)
                
                filtered.append(item)
                if is_dir: dir_count += 1
                else: file_count += 1
        
        # 如果数量超过限制，显示提示
        if len(filtered) > MAX_DISPLAY:
            lb.insert(tk.END, f"... (已隐藏 {len(filtered) - MAX_DISPLAY} 项，请使用搜索过滤) ...")
            lb.itemconfig(tk.END, {'fg': 'gray'})
            
        col_data["filtered_items"] = filtered
        col_data["status_var"].set(f"{dir_count} 文件夹, {file_count} 文件")


    def _sort_column(self, col_index, sort_by):
        """排序: sort_by = 'name' | 'date' | 'size'"""
        col_data = self.columns[col_index]
        items = col_data["all_items"]
        
        if sort_by == 'name':
            items.sort(key=lambda x: (not x[2], x[0].lower()))
        elif sort_by == 'date':
            items.sort(key=lambda x: (not x[2], x[4]), reverse=True) # 最新修改在前
        elif sort_by == 'size':
            items.sort(key=lambda x: (not x[2], x[3]), reverse=True) # 最大在前
            
        col_data["all_items"] = items
        # 重新应用过滤
        self._filter_items(col_index, col_data["search_var"].get().lower())

    def _on_key_left(self, col_index):
        """键盘左键：回到上一列"""
        if col_index > 0:
            prev_col = self.columns[col_index - 1]
            if prev_col["listbox"]:
                prev_col["listbox"].focus_set()

    def _on_key_right(self, col_index):
        """键盘右键：进入下一列"""
        # 1. 触发选中当前项
        self._on_column_select(col_index)
        
        # 2. 尝试聚焦下一列
        if col_index + 1 < len(self.columns):
            next_col = self.columns[col_index + 1]
            if next_col.get("is_preview", False):
                # 如果是预览列，不做聚焦操作，或者聚焦到文本框
                pass 
            elif next_col["listbox"]:
                next_col["listbox"].focus_set()
                # 默认选中第一项
                if next_col["listbox"].size() > 0:
                    next_col["listbox"].selection_clear(0, tk.END)
                    next_col["listbox"].selection_set(0)
                    next_col["listbox"].activate(0)

    def _open_file_location(self, col_index, item_index):
        """打开文件所在文件夹并选中文件"""
        try:
            col_data = self.columns[col_index]
            # 使用 filtered_items
            if item_index >= len(col_data["filtered_items"]):
                return
            _, full_path, _, _, _ = col_data["filtered_items"][item_index]
            
            full_path = os.path.abspath(full_path)
            
            if sys.platform == "win32":
                # explorer /select,"C:\path\to\file"
                subprocess.Popen(f'explorer /select,"{full_path}"')
            elif sys.platform == "darwin":
                subprocess.run(['open', '-R', full_path])
            else:
                # Linux/Unix
                parent = os.path.dirname(full_path)
                subprocess.run(['xdg-open', parent])
        except Exception as e:
            messagebox.showerror("错误", f"无法打开文件夹:\n{e}")

    def _reload_column(self, col_index):
        """刷新指定列的数据 (异步)"""
        if col_index >= len(self.columns): return
        
        col_data = self.columns[col_index]
        path = col_data["path"]
        lb = col_data["listbox"]
        
        # UI 立即反馈
        lb.delete(0, tk.END)
        lb.insert(tk.END, "Loading...")
        col_data["status_var"].set("正在加载...")
        
        # 记录这次加载的路径，防止快速切换导致数据错乱
        current_loading_path = path
        
        def _scan_task():
            try:
                all_data = []
                if not os.path.exists(path):
                    return

                # 使用 scandir 获取高性能列表
                with os.scandir(path) as it:
                    for entry in it:
                        name = entry.name
                        if name.startswith(".") or name in ["__pycache__", "node_modules", ".git", "$RECYCLE.BIN", "System Volume Information"]:
                            continue
                        full = entry.path
                        is_dir = entry.is_dir()
                        
                        # 获取大小和时间 (可能耗时)
                        size = 0
                        mtime = 0
                        try:
                            stat = entry.stat()
                            size = stat.st_size
                            mtime = stat.st_mtime
                        except: pass
                        
                        all_data.append((name, full, is_dir, size, mtime))
                
                # 默认排序
                all_data.sort(key=lambda x: (not x[2], x[0].lower()))
                
                # 回到主线程更新
                self.after(0, lambda: self._update_column_data(col_index, all_data, current_loading_path))
                
            except Exception as e:
                self.after(0, lambda: print(f"Scan failed: {e}"))

        threading.Thread(target=_scan_task, daemon=True).start()

    def _update_column_data(self, col_index, all_data, loaded_path):
        """更新列数据 (主线程)"""
        if col_index >= len(self.columns): return
        col_data = self.columns[col_index]
        
        # 确保路径一致 (防止快速切换)
        if col_data["path"] != loaded_path:
            return
            
        col_data["all_items"] = all_data
        
        # 重新应用当前的搜索过滤
        self._filter_items(col_index, col_data["search_var"].get())
        
        # 尝试自动调整宽度 (限制最大宽度)
        try:
            lb = col_data["listbox"]
            f = tkfont.Font(font=lb.cget("font"))
            max_w = 0
            # 只采样前 100 个
            for item in all_data[:100]:
                text = ("📂 " if item[2] else "📄 ") + item[0]
                w = f.measure(text)
                if w > max_w: max_w = w
            
            new_width = max(150, min(max_w + 40, 400))
            col_data["frame"].config(width=new_width)
        except: pass

    def _check_empty_click(self, event, lb):
        """拦截点击事件，防止点击空白处选中最后一行"""
        # 确保点击时获得焦点，以便响应 Ctrl+V 等快捷键
        lb.focus_set()
        
        # 获取最近的一行
        index = lb.nearest(event.y)
        # 获取该行的边界框
        bbox = lb.bbox(index)
        # 如果列表为空，或者点击位置在最后一行下方
        if bbox is None or event.y > bbox[1] + bbox[3]:
            return "break" # 阻止默认行为 (即不选中)

    def _on_backspace(self, event):
        """Backspace 返回上一级"""
        # 如果当前焦点在输入框，不拦截
        if isinstance(event.widget, tk.Entry) or isinstance(event.widget, tk.Text):
            return
            
        # 找到当前最深的一列
        if len(self.columns) > 1:
            # 移除最后一列
            self._clear_columns(len(self.columns) - 1)
            # 聚焦到新的一列
            if self.columns:
                last_col = self.columns[-1]
                if last_col.get("listbox"):
                    last_col["listbox"].focus_set()

    def _try_video_pause(self):
        """尝试暂停视频，成功返回 True"""
        if not self.columns:
            return False
            
        last_col = self.columns[-1]
        if last_col.get("is_preview", False):
            if "video_control" in last_col and "toggle_play" in last_col["video_control"]:
                try:
                    last_col["video_control"]["toggle_play"]()
                    return True
                except:
                    pass
        return False

    def _on_space_in_listbox(self, event):
        """处理列表中的空格键"""
        # 尝试处理视频暂停
        if self._try_video_pause():
            return "break"
        # 否则不做处理，让 Listbox 默认行为（选择）继续
        return None

    def _on_space(self, event):
        """Space 播放/暂停 (Global)"""
        # 如果焦点在输入框或按钮上，不拦截
        if isinstance(event.widget, (tk.Entry, ttk.Entry, tk.Text, tk.Button, ttk.Button)):
            return

        # 如果是 Listbox，已经在 _on_space_in_listbox 处理过
        if isinstance(event.widget, tk.Listbox):
            return

        self._try_video_pause()


    def _copy_file_with_progress(self, src, dst, callback):
        """带进度的文件复制"""
        try:
            with open(src, 'rb') as fsrc:
                with open(dst, 'wb') as fdst:
                    while True:
                        buf = fsrc.read(1024*1024) # 1MB chunks
                        if not buf: break
                        fdst.write(buf)
                        if callback:
                            callback(len(buf))
            shutil.copystat(src, dst)
        except Exception as e:
            print(f"Copy error: {src} -> {dst} : {e}")
            raise

    def _post_drop_refresh(self, col_index, drag_source_col_index, moved_count, copied_count):
        """拖放完成后的 UI 刷新"""
        try:
            # 刷新目标列
            self._reload_column(col_index)
            
            # 如果是内部移动，刷新源列
            if drag_source_col_index is not None:
                self._reload_column(drag_source_col_index)
                
            msg = []
            if moved_count: msg.append(f"移动 {moved_count} 个")
            if copied_count: msg.append(f"复制 {copied_count} 个")
            if msg:
                self.global_status_var.set(", ".join(msg) + " 完成")
            else:
                self.global_status_var.set("操作完成")
        except Exception as e:
            print(f"Post drop refresh error: {e}")

    def _process_drop_task(self, files, target_dir, is_internal, col_index, drag_source_col_index, progress_dlg):
        """后台线程处理复制/移动任务"""
        total_size = 0
        
        # 辅助函数：获取大小
        def get_size(path):
            total = 0
            try:
                if os.path.isfile(path):
                    total += os.path.getsize(path)
                elif os.path.isdir(path):
                    for r, d, f in os.walk(path):
                        for file in f:
                            total += os.path.getsize(os.path.join(r, file))
            except: pass
            return total

        # 1. 计算总大小
        self.after(0, lambda: progress_dlg.update(0, "正在计算文件大小..."))
        for f in files:
            total_size += get_size(f)
            
        current_bytes = 0
        moved_count = 0
        copied_count = 0
        
        # 进度更新回调
        def update_prog(inc_bytes, fname):
            nonlocal current_bytes
            current_bytes += inc_bytes
            pct = (current_bytes / total_size * 100) if total_size > 0 else 0
            # 使用 after 在主线程更新 UI (修复 lambda 闭包捕获问题)
            self.after(0, lambda p=pct, n=fname: progress_dlg.update(p, f"正在处理: {n}"))

        try:
            for src_path in files:
                if progress_dlg.cancelled:
                    self.after(0, lambda: messagebox.showinfo("取消", "操作已取消"))
                    break
                    
                # 检查文件是否存在
                if not os.path.exists(src_path): continue
                # 跳过自身
                if os.path.abspath(os.path.dirname(src_path)) == os.path.abspath(target_dir): continue

                filename = os.path.basename(src_path)
                dst_path = os.path.join(target_dir, filename)
                
                # 自动重命名
                if os.path.exists(dst_path):
                    base, ext = os.path.splitext(filename)
                    counter = 1
                    while os.path.exists(dst_path):
                        dst_path = os.path.join(target_dir, f"{base}_{counter}{ext}")
                        counter += 1
                
                # 移动逻辑
                if is_internal:
                    # 尝试快速移动 (os.rename)
                    try:
                        # 先获取大小用于更新进度
                        s = get_size(src_path)
                        os.rename(src_path, dst_path)
                        update_prog(s, filename)
                        moved_count += 1
                        continue
                    except OSError:
                        # 跨设备移动，回退到 Copy + Delete
                        pass

                # 复制逻辑 (或跨设备移动的复制阶段)
                if os.path.isdir(src_path):
                    # 递归复制文件夹
                    if not os.path.exists(dst_path):
                        os.makedirs(dst_path)
                    
                    for root, dirs, files_in_dir in os.walk(src_path):
                        # 创建目标子目录
                        rel_root = os.path.relpath(root, src_path)
                        dst_root = os.path.join(dst_path, rel_root)
                        for d in dirs:
                            dst_d = os.path.join(dst_root, d)
                            if not os.path.exists(dst_d):
                                os.makedirs(dst_d)
                        
                        for f in files_in_dir:
                            if progress_dlg.cancelled: break
                            src_f = os.path.join(root, f)
                            dst_f = os.path.join(dst_root, f)
                            self._copy_file_with_progress(src_f, dst_f, lambda b: update_prog(b, f))
                        
                        if progress_dlg.cancelled: break
                            
                    if is_internal: # 移动后的删除
                         shutil.rmtree(src_path)
                         moved_count += 1
                    else:
                         copied_count += 1
                         
                else:
                    # 单文件复制
                    self._copy_file_with_progress(src_path, dst_path, lambda b: update_prog(b, filename))
                    if is_internal:
                        os.remove(src_path)
                        moved_count += 1
                    else:
                        copied_count += 1

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("错误", str(e)))
        finally:
            # 稍微延迟关闭，确保用户能看到完成状态（尤其是小文件瞬间完成时）
            time.sleep(0.5)
            self.after(0, progress_dlg.close)
            self.after(0, lambda: self._post_drop_refresh(col_index, drag_source_col_index, moved_count, copied_count))

    def _on_drop_preview(self, event, current_preview_path):
        """处理文件拖入预览列"""
        try:
            # 预览列显示的是文件，所以目标目录应该是该文件所在的父目录
            target_dir = os.path.dirname(current_preview_path)
            
            files = self.tk.splitlist(event.data)
            
            # 找到对应的 col_index
            col_index = -1
            for i, col in enumerate(self.columns):
                if os.path.normpath(col.get("path")) == os.path.normpath(target_dir):
                    col_index = i
                    break
            
            pd = ProgressDialog(self, "正在导入文件...")
            threading.Thread(target=self._process_drop_task, 
                             args=(files, target_dir, False, col_index, None, pd),
                             daemon=True).start()

        except Exception as e:
            messagebox.showerror("拖放错误", f"无法导入文件:\n{e}")

    def _on_drop(self, event, col_index, files=None):
        """处理文件拖入"""
        try:
            col_data = self.columns[col_index]
            target_dir = col_data["path"]
            
            if files is None:
                files = self.tk.splitlist(event.data)
            
            # 判断是否内部拖拽
            is_internal = getattr(self, "_is_internal_drag", False)
            drag_source_col_index = getattr(self, "_drag_source_col_index", None)
            
            # 创建进度条窗口
            pd = ProgressDialog(self, "正在处理文件...")
            
            # 启动线程
            threading.Thread(target=self._process_drop_task, 
                             args=(files, target_dir, is_internal, col_index, drag_source_col_index, pd),
                             daemon=True).start()
                             
        except Exception as e:
            messagebox.showerror("拖放错误", f"无法导入文件:\n{e}")

    def _on_drag_init(self, event, col_index):
        """拖拽开始"""
        try:
            col_data = self.columns[col_index]
            sel = col_data["listbox"].curselection()
            if not sel: return
            
            # 获取选中的文件路径
            files = []
            for i in sel:
                if i < len(col_data["filtered_items"]):
                    _, full_path, _, _, _ = col_data["filtered_items"][i]
                    files.append(os.path.normpath(full_path))
            
            if not files: return
            
            self._is_internal_drag = True
            self._drag_source_col_index = col_index
            
            # 生成 Tcl 列表格式的数据
            data = self.tk.call('list', *files)
            
            # 返回 (action, type, data)
            return ('move', DND_FILES, data)
        except Exception as e:
            print(f"Drag init error: {e}")

    def _on_drag_end(self, event):
        """拖拽结束"""
        self._is_internal_drag = False
        self._drag_source_col_index = None

    def _on_copy(self, event):
        """Ctrl+C 快捷键"""
        lb = event.widget
        # 查找对应的 col_index
        for i, col in enumerate(self.columns):
            if col["listbox"] == lb:
                if lb.curselection():
                    self._perform_copy(i)
                break

    def _on_paste(self, event):
        """Ctrl+V 快捷键"""
        lb = event.widget
        # 查找对应的 col_index
        for i, col in enumerate(self.columns):
            if col["listbox"] == lb:
                self._perform_paste(i)
                break
                
    def _on_delete(self, event):
        """Delete 快捷键"""
        lb = event.widget
        for i, col in enumerate(self.columns):
            if col["listbox"] == lb:
                if lb.curselection():
                    self._perform_delete(i)
                break

    def _perform_copy(self, col_index, item_index=None):
        """执行复制 (支持多选)"""
        col_data = self.columns[col_index]
        lb = col_data["listbox"]
        sel = lb.curselection()
        
        if not sel: return
        
        files_to_copy = []
        for idx in sel:
            if idx < len(col_data["filtered_items"]):
                files_to_copy.append(col_data["filtered_items"][idx][1]) # full_path
        
        if not files_to_copy: return

        self._clipboard_op = 'copy'
        self._clipboard_files = files_to_copy
        
        # 同步到系统剪贴板
        ClipboardUtils.set_files(files_to_copy)
        
        self.global_status_var.set(f"已复制 {len(files_to_copy)} 个项目")

    def _perform_cut(self, col_index, item_index=None):
        """执行剪切 (支持多选)"""
        col_data = self.columns[col_index]
        lb = col_data["listbox"]
        sel = lb.curselection()
        
        if not sel: return

        files_to_cut = []
        for idx in sel:
            if idx < len(col_data["filtered_items"]):
                files_to_cut.append(col_data["filtered_items"][idx][1]) # full_path

        if not files_to_cut: return
        
        self._clipboard_op = 'cut'
        self._clipboard_files = files_to_cut
        
        # 同步到系统剪贴板
        ClipboardUtils.set_files(files_to_cut)
        
        self.global_status_var.set(f"已剪切 {len(files_to_cut)} 个项目")

    def _perform_paste(self, col_index):
        col_data = self.columns[col_index]
        target_dir = col_data["path"]
        
        # 1. 优先获取系统剪贴板文件
        sys_files = ClipboardUtils.get_files()
        
        source_files = []
        op = 'copy'
        
        if sys_files:
            source_files = sys_files
            # 检查是否与内部剪贴板一致 (判断是否为剪切操作)
            if self._clipboard_files and set(self._clipboard_files) == set(sys_files):
                if self._clipboard_op == 'cut':
                    op = 'cut'
        elif self._clipboard_files:
            # 如果系统剪贴板无文件，尝试使用内部缓存
            source_files = self._clipboard_files
            op = self._clipboard_op or 'copy'
            
        if not source_files:
            return
        
        # 预检查和确认 (在主线程完成交互)
        final_files = []
        for src_path in source_files:
            if not os.path.exists(src_path): continue
            
            # 简单检查是否复制到自身子目录
            try:
                if os.path.abspath(target_dir).startswith(os.path.abspath(src_path) + os.sep):
                    continue
            except:
                pass
            
            final_files.append(src_path)
            
        if not final_files: return

        # 准备参数
        is_internal = (op == 'cut')
        
        # 创建进度条
        pd = ProgressDialog(self, "正在粘贴...")
        
        # 启动后台线程
        threading.Thread(target=self._process_drop_task, 
                         args=(final_files, target_dir, is_internal, col_index, None, pd),
                         daemon=True).start()
        
        # 如果是剪切，清空内部剪贴板状态
        if op == 'cut':
            self._clipboard_files = []
            self._clipboard_op = None

    def _perform_delete(self, col_index, item_index=None):
        """执行删除 (支持多选)"""
        col_data = self.columns[col_index]
        lb = col_data["listbox"]
        sel = lb.curselection()
        
        if not sel: return
        
        items_to_delete = []
        for idx in sel:
            if idx < len(col_data["filtered_items"]):
                items_to_delete.append(col_data["filtered_items"][idx])
        
        if not items_to_delete: return
        
        msg = f"确定要永久删除这 {len(items_to_delete)} 个项目吗？" if len(items_to_delete) > 1 else f"确定要永久删除 '{items_to_delete[0][0]}' 吗？"
        
        if messagebox.askyesno("删除确认", msg):
            success_count = 0
            for name, full_path, _, _, _ in items_to_delete:
                try:
                    if os.path.isdir(full_path):
                        shutil.rmtree(full_path)
                    else:
                        os.remove(full_path)
                    success_count += 1
                except Exception as e:
                    print(f"Delete failed for {name}: {e}")
            
            if success_count > 0:
                self._reload_column(col_index)
                self._clear_columns(col_index + 1)
                self.global_status_var.set(f"已删除 {success_count} 个项目")
            else:
                 messagebox.showerror("错误", "删除失败，请检查权限")

    def _rename_item(self, col_index, item_index):
        """重命名文件/文件夹"""
        col_data = self.columns[col_index]
        if item_index >= len(col_data["filtered_items"]): return
        
        name, full_path, _, _, _ = col_data["filtered_items"][item_index]
        
        new_name = simpledialog.askstring("重命名", "请输入新名称:", initialvalue=name, parent=self)
        if new_name and new_name != name:
            try:
                new_path = os.path.join(os.path.dirname(full_path), new_name)
                os.rename(full_path, new_path)
                
                # 刷新当前列
                self._reload_column(col_index)
                # 清除后续列，因为路径可能变了
                self._clear_columns(col_index + 1)
            except Exception as e:
                messagebox.showerror("错误", f"重命名失败:\n{e}")

    def _new_folder(self, col_index):
        """新建文件夹"""
        col_data = self.columns[col_index]
        path = col_data["path"]
        
        new_name = simpledialog.askstring("新建文件夹", "请输入文件夹名称:", parent=self)
        if new_name:
            try:
                new_path = os.path.join(path, new_name)
                os.mkdir(new_path)
                self._reload_column(col_index)
            except Exception as e:
                messagebox.showerror("错误", f"创建失败:\n{e}")

    # 下面是废弃的方法，但为了避免报错可以保留为空或删除
    # 为保持 diff 简洁，这里我们直接替换掉原来的事件处理方法，不需要保留旧的

if __name__ == "__main__":
    app = FileViewerApp()
    app.mainloop()